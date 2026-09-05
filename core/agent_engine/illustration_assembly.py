"""G-7 收尾：AI 配图增量装配（备用状态）。

目标态：用户在 .env / 设置页填入图片 API（BMP_IMAGE_PROVIDER + BMP_IMAGE_BASE_URL +
BMP_IMAGE_MODEL + BMP_IMAGE_API_KEY）并把 BMP_ILLUSTRATION_ENABLED 置 true，即端到端出图：
读章节 citation_ledger["_illustrations"] 建议 → 调 AiImageSkill 生成 → 在正文对应
[插图位置] 标记处插入图片引用（Markdown 图片语法；HTML 导出经 markdown 渲染为 <img>，
DOCX 导出用 embed_images_into_docx 嵌入）。开关关=现状（建议留元数据，正文无脚手架）。

供应商取舍：自定义 provider 走 OpenAI 兼容 images/generations（事实标准，响应 schema
统一）；无 key / 服务失败统一降级 fallback（可离线占位图），不阻断章节生成。
"""

from __future__ import annotations

import base64
import logging
import re

logger = logging.getLogger(__name__)

# 与 illustration_guard 同源的正文标记（装配消费的就是尚未剥离的 gate 文本）
_MARKER_RE = re.compile(r"\[插图位置[：:]?([^\]]*)\]")

_MAX_IMAGES_PER_CHAPTER = 3
_MAX_B64_CHARS = 4_000_000  # 单图 base64 上限（约 3MB），防 ledger 膨胀


def _illustration_markdown(image: dict, index: int) -> str:
    """Return one stable Markdown reference for a persisted illustration."""
    src = str(image.get("src") or "").strip()
    if not src:
        return ""
    alt = str(image.get("alt") or f"配图{index}").strip()
    return f"![{alt}]({src})"


def attach_stored_illustrations(content: str, images: list[dict] | None) -> str:
    """Attach persisted images without invoking an image service."""
    text = str(content or "")
    valid = [img for img in (images or []) if isinstance(img, dict) and img.get("src")]
    if not valid or all(str(img.get("src")) in text for img in valid):
        return text
    markers = list(_MARKER_RE.finditer(text))
    refs = [_illustration_markdown(img, i) for i, img in enumerate(valid, start=1)]
    refs = [ref for ref in refs if ref]
    if not refs:
        return text
    if markers:
        out: list[str] = []
        last = 0
        consumed = min(len(markers), len(refs))
        for marker, ref in zip(markers[:consumed], refs[:consumed]):
            out.append(text[last:marker.start()])
            out.append(ref)
            last = marker.end()
        out.append(text[last:])
        result = "".join(out)
        extra = refs[consumed:]
    else:
        result = text
        extra = refs
    if extra:
        result = result.rstrip() + "\n\n" + "\n\n".join(extra) + "\n"
    return result


def get_illustration_params(force_enabled: bool = False, provider_override: str | None = None) -> dict | None:
    """按 settings 解析装配参数。返回 None=不装配（开关关或可用性不满足）。"""
    from core.settings import get_settings

    s = get_settings()
    if not force_enabled and not s.illustration_enabled:
        return None
    provider = str(provider_override or s.image_provider or "fallback")
    if provider == "default":
        provider = "fallback"
    if provider == "fallback":
        key_ok = True  # fallback（含离线占位）无需 key
    else:
        import os

        env_key = ""
        if provider == "volcengine":
            env_key = os.getenv("VOLCENGINE_API_KEY", "")
        elif provider == "google":
            env_key = os.getenv("GOOGLE_API_KEY", "")
        key_ok = bool(env_key or s.image_api_key)
    if not key_ok:
        return None
    return {
        "provider": provider,
        "base_url": str(s.image_base_url or ""),
        "model": str(s.image_model or ""),
        "api_key": "" if provider == "fallback" else str(s.image_api_key or ""),
    }


async def assemble_illustrations(
    project_id: str,
    content: str,
    suggestions: list[str],
    params: dict | None = None,
    image_size: str = "landscape_16_9",
) -> tuple[str, list[dict], str]:
    """把配图建议装配进正文。

    - content 含 [插图位置…] 标记（gate 文本）：逐标记生成图片并原位替换为 Markdown 图片引用；
    - content 无标记但 suggestions 非空：图片引用追加到章末（保序，不散插入正文）；
    - params 为 None（开关关/无 key）：原样返回，正文不动（现状语义）。

    返回 (新正文, images 元数据列表, 状态说明)。
    """
    text = str(content or "")
    if params is None:
        return text, [], "disabled"

    markers = list(_MARKER_RE.finditer(text))
    targets: list[str] = [m.group(0) for m in markers[:_MAX_IMAGES_PER_CHAPTER]]
    if not targets and suggestions:
        targets = [str(s) for s in suggestions[:_MAX_IMAGES_PER_CHAPTER]]
    if not targets:
        return text, [], "no_suggestions"

    from core.skill_engine.base import SkillContext
    from services.generate.skills.ai_image_skill import AiImageSkill

    skill = AiImageSkill()
    images: list[dict] = []
    for i, desc in enumerate(targets, start=1):
        prompt = re.sub(r"[\[\]【】\*]", "", desc).strip() or "配图"
        ctx = SkillContext(
            project_id=str(project_id or ""),
            db=None,
            llm=None,
            parameters={
                "prompt": prompt,
                "provider": params.get("provider") or "fallback",
                "image_size": image_size,
                "image_base_url": params.get("base_url") or "",
                "image_model": params.get("model") or "",
                "api_key": params.get("api_key") or "",
            },
        )
        result = await skill.safe_execute(ctx)
        if not result.success or not result.data:
            logger.warning("配图生成失败（跳过该标记）: %s", result.error)
            continue
        data = result.data
        src = ""
        if data.get("image_url"):
            src = str(data["image_url"])
        elif data.get("base64"):
            b64 = str(data["base64"])
            if len(b64) > _MAX_B64_CHARS:
                logger.warning("配图 base64 超限（%d 字符），跳过", len(b64))
                continue
            src = f"data:image/png;base64,{b64}"
        if not src:
            continue
        alt = f"配图{i}：{prompt[:60]}"
        images.append(
            {
                "index": i,
                "alt": alt,
                "src": src,
                "provider": data.get("provider", ""),
                "placeholder": bool(data.get("placeholder", False)),
                "marker": desc,
            }
        )

    if not images:
        return text, [], "generation_failed"

    # 原位替换：按 markers 命中顺序消费 images；多余 images 追加到章末
    out_parts: list[str] = []
    last = 0
    img_iter = iter(images)
    for m in markers[:_MAX_IMAGES_PER_CHAPTER]:
        img = next(img_iter, None)
        if img is None:
            break
        out_parts.append(text[last : m.start()])
        out_parts.append(f"![{img['alt']}]({img['src']})")
        last = m.end()
    if last:
        out_parts.append(text[last:])
        new_text = "".join(out_parts)
    else:
        new_text = text
    # 无标记、仅建议：追加到章末
    consumed = sum(1 for _ in markers[:_MAX_IMAGES_PER_CHAPTER])
    extra = images[consumed:] if markers else images
    if not markers:
        appendix = "\n\n".join(f"![{img['alt']}]({img['src']})" for img in extra)
        if appendix:
            new_text = new_text.rstrip() + "\n\n" + appendix + "\n"
    elif extra:
        appendix = "\n\n".join(f"![{img['alt']}]({img['src']})" for img in extra)
        new_text = new_text.rstrip() + "\n\n" + appendix + "\n"

    return new_text, images, "assembled"


def _embed_images_into_docx_legacy(docx_path: str) -> int:
    """把正文中 Markdown 图片引用（含 data URI）嵌入 DOCX：图片段落 add_picture。

    供导出脚本在生成 docx 后调用；返回嵌入图片数。data URI 过大时缩到不超限。
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx 不可用，跳过 DOCX 配图嵌入")
        return 0
    import io as _io

    doc = DocxDocument(docx_path)
    count = 0
    for para in list(doc.paragraphs):
        text = para.text or ""
        m = re.search(r"!\[([^\]]*)\]\((data:image/[^;]+;base64,([^)]+)|https?://[^)]+)\)", text)
        if not m:
            continue
        src = m.group(2)
        b64 = m.group(3) if src.startswith("data:") else ""
        url = src if src.startswith("http") else ""
        try:
            img_bytes = base64.b64decode(b64) if b64 else None
            if img_bytes is None:
                import httpx

                resp = httpx.get(url, timeout=60, follow_redirects=True)
                resp.raise_for_status()
                img_bytes = resp.content
            # 清空原段文字，插入图片
            for run in list(para.runs):
                run.text = ""
            stream = _io.BytesIO(img_bytes)
            para.add_run().add_picture(stream)
            count += 1
        except Exception as e:  # noqa: BLE001
            logger.warning("DOCX 配图嵌入失败（保留引用文本）: %s", e)
            continue
    if count:
        doc.save(docx_path)
    return count


def embed_images_into_docx(docx_path: str) -> int:
    """Embed persisted Markdown image references, including local base64 refs."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx 不可用，跳过 DOCX 配图嵌入")
        return 0
    import io as _io

    image_re = re.compile(
        r"!\[([^\]]*)\]\(((?:\[image omitted\][^)]+|data:[^;]+;base64,[^)]+|https?://[^)]+))\)"
    )
    doc = DocxDocument(docx_path)
    count = 0
    for para in list(doc.paragraphs):
        match = image_re.search(para.text or "")
        if not match:
            continue
        src = match.group(2)
        if src.startswith("[image omitted]"):
            b64, url = src[len("[image omitted]"):], ""
        elif src.startswith("data:"):
            b64, url = src.split(",", 1)[1], ""
        else:
            b64, url = "", src
        try:
            img_bytes = base64.b64decode(b64) if b64 else None
            if img_bytes is None:
                import httpx
                response = httpx.get(url, timeout=60, follow_redirects=True)
                response.raise_for_status()
                img_bytes = response.content
            for run in list(para.runs):
                run.text = ""
            para.add_run().add_picture(_io.BytesIO(img_bytes))
            count += 1
        except Exception as exc:  # noqa: BLE001
            logger.warning("DOCX 配图嵌入失败（保留引用文本）: %s", exc)
    if count:
        doc.save(docx_path)
    return count


def illustration_img_tags(images: list[dict]) -> list[str]:
    """HTML 导出用：<img> 标签列表（手工拼 HTML 的导出方使用）。"""
    return [f'<img src="{img["src"]}" alt="{img["alt"]}" />' for img in images]


__all__ = [
    "get_illustration_params",
    "assemble_illustrations",
    "embed_images_into_docx",
    "illustration_img_tags",
]
