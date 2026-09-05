from __future__ import annotations

from core.doc_engine.parsers.base import ParsedDocument


class DocxParser:
    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document

        doc = Document(file_path)
        text_parts = []
        tables = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(
                {
                    "index": i,
                    "headers": rows[0] if rows else [],
                    "rows": rows[1:] if len(rows) > 1 else [],
                }
            )

        return ParsedDocument(
            text="\n".join(text_parts),
            tables=tables,
            images=[],
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "parser": "python-docx",
            },
        )
