"""结构化流水线（P-A 挂接点）：在现有解析链完成后追加版面/表格/分块/实体产物。

- 纯计算 + JSON 落盘（data/structured/<key>/），实体另入 DB（tender_entities）；
- 不改变现有产物格式与消费方；任何异常由调用方捕获后静默降级（零回归）；
- 服务端挂接走后台线程（不阻塞解析接口）。
"""

from __future__ import annotations

import hashlib
import json
import logging
from datetime import datetime
from pathlib import Path

from core.doc_engine.chunker import ChunkConfig, chunk_document
from core.doc_engine.layout import LayoutBlock, build_layout
from core.doc_engine.table_structure import structure_tables

logger = logging.getLogger(__name__)

STRUCTURED_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "structured"


def _artifact_key(file_path: str, project_id: str) -> str:
    if project_id:
        return project_id[:8] or project_id
    return hashlib.sha256(Path(file_path).read_bytes()[:65536]).hexdigest()[:12]


def _extract_pages(file_path: str) -> list[str]:
    """按页文本（页码证据用）；非 PDF 返回单元素列表。"""
    path_obj = Path(file_path)
    suffix = path_obj.suffix.lower()
    if suffix == ".pdf":
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            return [p.extract_text() or "" for p in pdf.pages]
    if suffix == ".docx":
        from docx import Document

        doc = Document(file_path)
        return ["\n".join(p.text for p in doc.paragraphs if p.text.strip())]
    return [path_obj.read_text(encoding="utf-8", errors="ignore")]


def run_structuring(
    file_path: str,
    project_id: str = "",
    document_id: str = "",
    chunk_config: ChunkConfig | None = None,
    llm_gateway=None,
    use_llm_entities: bool = True,
    out_root: Path | None = None,
    save_db: bool = False,
) -> dict:
    """完整结构化：版面 → 表格 → 分块 → 实体，产物 JSON 落盘，返回 summary。"""
    from core.doc_engine.entity_extractor import extract_entities

    cfg = chunk_config or ChunkConfig()
    key = _artifact_key(file_path, project_id)
    out_dir = (out_root or STRUCTURED_DIR) / key
    out_dir.mkdir(parents=True, exist_ok=True)

    # 1) 版面结构化
    layout = build_layout(file_path)
    blocks: list[LayoutBlock] = layout.blocks
    blocks_dicts = [b.to_dict() for b in blocks]

    # 2) 表格结构化（平面表来自现有解析器，叠加不替换）
    from core.doc_engine import get_parser

    parsed = get_parser(Path(file_path).suffix).parse(file_path)
    tables = structure_tables(parsed.tables, blocks_dicts)

    # 3) 条款级语义分块
    chunked = chunk_document(parsed.text, blocks=blocks, tables=tables, project_id=project_id or key, config=cfg)

    # 4) 实体抽取双路（页码证据来自按页文本）
    pages = _extract_pages(file_path)
    entities = extract_entities(parsed.text, pages=pages, llm_gateway=llm_gateway, use_llm=use_llm_entities)

    # 落盘
    artifacts: dict[str, Path] = {}
    (out_dir / "layout.json").write_text(
        json.dumps(
            {**layout.to_dict(), "file": file_path, "generated_at": datetime.now().isoformat(timespec="seconds")},
            ensure_ascii=False,
            indent=1,
        ),
        encoding="utf-8",
    )
    artifacts["layout"] = out_dir / "layout.json"
    (out_dir / "tables.json").write_text(
        json.dumps(
            {
                "file": file_path,
                "generated_at": datetime.now().isoformat(timespec="seconds"),
                "table_count": len(tables),
                "cross_page_merged": sum(1 for t in tables if t.merged_from_pages),
                "tables": [t.to_dict() for t in tables],
                "markdown": [t.to_markdown() for t in tables],
            },
            ensure_ascii=False,
            indent=1,
        ),
        encoding="utf-8",
    )
    artifacts["tables"] = out_dir / "tables.json"
    (out_dir / "chunks.json").write_text(
        json.dumps(
            {**chunked, "file": file_path, "generated_at": datetime.now().isoformat(timespec="seconds")},
            ensure_ascii=False,
            indent=1,
        ),
        encoding="utf-8",
    )
    artifacts["chunks"] = out_dir / "chunks.json"
    (out_dir / "entities.json").write_text(
        json.dumps(
            {**entities, "file": file_path, "generated_at": datetime.now().isoformat(timespec="seconds")},
            ensure_ascii=False,
            indent=1,
        ),
        encoding="utf-8",
    )
    artifacts["entities"] = out_dir / "entities.json"

    summary = {
        "file": file_path,
        "project_id": project_id,
        "document_id": document_id,
        "artifact_key": key,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "layout": {"block_count": len(blocks), "parser_used": layout.parser_used, "degraded": layout.degraded},
        "tables": {
            "table_count": len(tables),
            "cross_page_merged": sum(1 for t in tables if t.merged_from_pages),
            "two_level_headers": sum(1 for t in tables if t.header_level >= 2),
        },
        "chunks": chunked["stats"],
        "entities": entities["stats"],
        "paths": {k: str(v) for k, v in artifacts.items()},
    }
    (out_dir / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=1), encoding="utf-8")

    if save_db and project_id:
        _persist_db(project_id, document_id, entities["entities"], artifacts, summary)

    return summary


def _persist_db(project_id: str, document_id: str, entities: list[dict], artifacts: dict, summary: dict) -> None:
    """实体入库 + 产物登记（独立 DB 会话；失败仅告警）。"""
    try:
        import asyncio

        from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

        from core.settings import get_settings
        from services.models import StructuredArtifact, TenderEntity

        async def _run() -> None:
            # 独立引擎（可能运行在后台线程自有事件循环，严禁复用服务端全局引擎，
            # 否则 asyncpg 连接绑定旧循环导致 "another operation is in progress"）
            engine = create_async_engine(get_settings().get_database_url())
            factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
            try:
                async with factory() as db:
                    for e in entities:
                        db.add(
                            TenderEntity(
                                project_id=project_id or None,
                                document_id=document_id or None,
                                entity_type=e["entity_type"],
                                value=e["value"][:500],
                                norm=(e.get("norm") or "")[:500],
                                source=e["source"],
                                confidence=float(e.get("confidence") or 0),
                                page=int(e.get("page") or 0),
                                evidence=e.get("evidence") or "",
                                conflict=bool(e.get("conflict")),
                                review_status=e.get("review_status") or "auto",
                            )
                        )
                    for atype, path in artifacts.items():
                        db.add(
                            StructuredArtifact(
                                project_id=project_id or None,
                                document_id=document_id or None,
                                artifact_type=atype,
                                path=str(path),
                                summary=summary.get(atype, {}),
                            )
                        )
                    await db.commit()
            finally:
                await engine.dispose()

        asyncio.run(_run())
    except Exception as exc:  # noqa: BLE001
        logger.warning("结构化产物入库失败（不影响解析主链路）: %s", exc)


def run_structuring_background(file_path: str, project_id: str = "", document_id: str = "", **kwargs) -> None:
    """后台线程执行（服务端挂接用；异常只记日志，零回归）。"""
    import threading

    def _worker() -> None:
        try:
            kwargs.setdefault("save_db", True)
            summary = run_structuring(file_path, project_id=project_id, document_id=document_id, **kwargs)
            logger.info(
                "P-A 结构化完成 project=%s: chunks=%s tables=%s entities=%s",
                project_id,
                summary["chunks"]["total_chunks"],
                summary["tables"]["table_count"],
                summary["entities"]["total"],
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("P-A 结构化失败（降级，不影响主链路）: %s", exc)

    threading.Thread(target=_worker, name="pa-structuring", daemon=True).start()
