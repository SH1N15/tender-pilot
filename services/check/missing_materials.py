"""Worker I 任务2：《需补充材料清单》确定性提取与渲染（无 LLM）。

场景价值：全量检查后，把"文本重写无法根治、必须线下收集材料"的缺口
（签字盖章、证书扫描件、最终报价、检测报告等）汇成一张可执行的收集清单，
作为用户"放弃还是补料"决策的依据。

数据来源（两路合并，按 (chapter_id, 指纹) 去重）：
1. 检查报告 results —— fail/warning 的 fact 型缺失 finding
   （文本命中 待补充/知识库无据/未提供/缺少/缺失/缺 类关键词）；
2. 章节正文 —— 证据硬门拒答标记【待补充】(原因：…；待补充：…) 与
   （知识库无据，待补充…）占位语，连同上下文定位到章节位置。

行结构：{position, chapter_id, detail, missing_type, suggested_material,
priority, source}；priority ∈ high/medium/low（确定性关键词规则）。
"""

from __future__ import annotations

import io
import re

# 事实型缺失关键词（finding 文本命中即入清单；"缺输入:" 为执行 skip 原因，不在此类）
MISSING_RE = re.compile(r"待补充|知识库无据|未提供|缺少|缺失|缺如|未响应|无法核实|需补充")
_NOT_REQUIRED_RE = re.compile(
    r"无需|无须|不需要|不必|不用|免收|免于|未要求|不要求|不适用|视为合规|无需补充|无需提交|无需提供|无需上传"
)
_LATER_STAGE_RE = re.compile(
    r"中标后|成交后|合同签订后|签约后|履约阶段|后续履约|后续执行|后续办理|实施阶段|交付阶段|递交后|提交后"
)
_CURRENT_STAGE_RE = re.compile(r"当前阶段|本阶段|投标阶段|资格审查阶段|递交阶段")
_FINAL_ACTION_RE = re.compile(
    r"纸质|密封袋|密封标识|骑缝章|正本副本|物理封装|最终PDF|最终文件|平台客户端|加密上传|开标解密|解密测试|签到|电子签章覆盖|签章覆盖|平台回执|最终提交|签章时间戳|CA证书|数字证书"
)
_SATISFIED_RE = re.compile(
    r"(?:已(?:经)?(?:全部|基本|完整)?涵盖|均已(?:提供|覆盖)|材料(?:已)?齐全|"
    r"(?:未发现|未见)(?:明显)?(?:的)?(?:必传文件|材料|资料|文件)?(?:缺失|遗漏|不足|负偏离)|"
    r"无(?:明显)?(?:的)?(?:必传文件|材料|资料|文件)?(?:缺失|遗漏|不足))"
)

# “当前阶段无需提交、后续阶段再办理”是流程事项，不应被归为企业缺料。
# 仅依赖通用状态语义，不绑定某一种材料、金额或项目文本。
def is_non_current_material_matter(text: str) -> bool:
    value = re.sub(r"[\s、，。；;：:（）()【】[\]“”'\"`]", "", str(text or ""))
    not_required = bool(_NOT_REQUIRED_RE.search(value))
    later_stage = bool(_LATER_STAGE_RE.search(value))
    current_stage = bool(_CURRENT_STAGE_RE.search(value))
    # Final-file assembly and platform/physical signing are workflow actions,
    # not enterprise evidence gaps.  Keep them visible as human follow-ups.
    if _FINAL_ACTION_RE.search(value) and re.search(
        r"无法|未能|未提供|缺少|需确认|请确认|无法核实|尚未|最终", value
    ):
        return True
    # Text-only signature findings describe a final artifact that a human
    # must sign/seal or visually verify. They are workflow actions, not a
    # request to upload another enterprise fact, when the finding explicitly
    # says that the rendered signature cannot be verified.
    if re.search(r"签字|签章|盖章", value) and re.search(
        r"实际签署|实际盖章|签署痕迹|无法直接验证|签章截图|验证报告|模板或草稿", value
    ):
        return True
    return (not_required and (later_stage or current_stage)) or bool(
        re.search(r"(?:无需|无须|不适用|视为合规).{0,80}(?:材料|文件|证明|提交|上传)", value)
    )

# 章节正文占位标记：evidence_gate 的 【待补充】(原因：R；待补充：V)（ASCII/全角括号，允许裸标记）
PEND_MARKER_RE = re.compile(r"【待补充】(?:[（(]([^）)]*)[）)])?")
KB_MISSING_SENTENCE_RE = re.compile(r"（知识库无据，待补充[^）]*）?[^。\n]*")

# check_id → 缺失要件类型（报告 findings 侧的分类标签）
_CHECK_TYPE_LABELS = {
    "qualification_check": "资质证书",
    "signature_check": "签章/签字",
    "deposit_check": "保证金凭证",
    "sample_report_check": "检测/样品报告",
    "pricing_check": "报价数据",
    "pricing_logic_check": "报价数据",
    "validity_check": "有效期证明",
    "doc_integrity_check": "附件/文件",
    "joint_bid_check": "联合体协议",
    "mandatory_req_check": "★▲参数响应证明",
    "cross_check": "跨章节事实数据",
    "consistency_check": "跨章节事实数据",
}

# 要件类型推断（优先级从高到低，首个命中）
_TYPE_RULES: list[tuple[re.Pattern, str]] = [
    (re.compile(r"报价|金额|价格|限价|大写"), "报价数据"),
    (re.compile(r"公章|盖章|签章|签字|签名|骑缝|法人"), "签章/签字"),
    (re.compile(r"证书|资质|执照|许可"), "资质证书"),
    (re.compile(r"保证金|保函"), "保证金凭证"),
    (re.compile(r"检测|样品|CMA|CNAS|cnas|cma"), "检测/样品报告"),
    (re.compile(r"业绩|合同|中标通知|验收"), "业绩证明"),
    (re.compile(r"授权|委托"), "授权委托文件"),
    (re.compile(r"日期|工期|质保|有效期"), "日期/期限数据"),
]

# 建议收集材料（首个命中）
_MATERIAL_RULES: list[tuple[re.Pattern, str]] = [
    (re.compile(r"报价|金额|价格|大写"), "最终报价单/分项报价表（法定代表人签字+加盖公章）"),
    (re.compile(r"公章|盖章|签章|签字|签名|骑缝|法人"), "签字盖章页扫描件（法人签字+公司公章，含骑缝章）"),
    (re.compile(r"证书|资质|执照|许可"), "资质证书/营业执照扫描件（须在有效期内）"),
    (re.compile(r"保证金|保函"), "保证金缴纳凭证或银行保函扫描件"),
    (re.compile(r"检测|样品|CMA|CNAS|cnas|cma"), "CMA/CNAS 检测报告或样品证明材料扫描件"),
    (re.compile(r"业绩|合同|中标通知|验收"), "同类项目业绩合同/中标通知书/验收证明扫描件"),
    (re.compile(r"授权|委托"), "法定代表人授权委托书（签字盖章原件扫描）"),
    (re.compile(r"日期|工期|质保|有效期"), "明确日期/工期数值及其依据文件（上传企业私有库）"),
]
_DEFAULT_MATERIAL = "相关佐证材料（整理后上传企业私有知识库，再触发重新检查修复）"

_HIGH_RE = re.compile(r"废标|资格|资质|证书|盖章|签章|签字|公章|原件|保证金|报价|法定代表人|授权")
_MEDIUM_RE = re.compile(r"检测|样品|业绩|合同|验收|证明|扫描|有效期|保函|委托")

_PRIORITY_ORDER = {"high": 0, "medium": 1, "low": 2}


def is_missing_material_text(text: str) -> bool:
    """确定性判定：一段 finding/正文文本是否属于"缺料"类事实缺失。"""
    t = str(text or "")
    if not t:
        return False
    if t.startswith("缺输入"):  # 图执行 skip 的基础设施原因，不算缺料
        return False
    if re.search(r"未见(?:明显)?(?:问题|重复计价|异常|缺陷|冲突)", t):
        return False
    # Positive coverage statements must not become upload tasks merely because
    # they contain the characters “缺失/未发现”. Genuine “未提供 X” findings
    # still pass through this guard and remain actionable.
    if _SATISFIED_RE.search(t) and not re.search(
        r"(?:待补充|需补充|缺少[^，。；;\n]{0,40}|未提供[^，。；;\n]{0,40})", t
    ):
        return False
    return bool(MISSING_RE.search(t))


def classify_missing_type(text: str, check_id: str = "", check_name: str = "") -> str:
    if is_non_current_material_matter(text):
        return "后续履约/人工事项"
    for pat, label in _TYPE_RULES:
        if pat.search(text or ""):
            return label
    return _CHECK_TYPE_LABELS.get(check_id) or check_name or "事实数据"


def suggest_material(text: str) -> str:
    if is_non_current_material_matter(text):
        return "当前阶段无需上传材料；请在招标文件规定的适用阶段执行"
    for pat, hint in _MATERIAL_RULES:
        if pat.search(text or ""):
            return hint
    return _DEFAULT_MATERIAL


def priority_of(text: str) -> str:
    if _HIGH_RE.search(text or ""):
        return "high"
    if _MEDIUM_RE.search(text or ""):
        return "medium"
    return "low"


def reconcile_missing_findings(findings: list[dict], feedback: dict | None) -> list[dict]:
    """Remove findings proven resolved by the focused repair/recheck loop."""
    tasks = (feedback or {}).get("tasks") if isinstance(feedback, dict) else []
    if not isinstance(tasks, list) or not tasks:
        return list(findings or [])
    out: list[dict] = []
    for finding in findings or []:
        check_id = str(finding.get("check_id") or "")
        chapter_id = str(finding.get("chapter_id") or "")
        matched = [
            task for task in tasks
            if isinstance(task, dict)
            and str(task.get("check_id") or "") == check_id
            and (not chapter_id or not task.get("chapter_id") or str(task.get("chapter_id")) == chapter_id)
        ]
        resolved = any(
            bool(task.get("fixed"))
            or str(task.get("status_after") or "").lower() == "pass"
            or str((task.get("recheck") or {}).get("status") or "").lower() == "pass"
            for task in matched
        )
        if not resolved:
            out.append(finding)
    return out


def build_action_summary(
    check_results: list[dict] | None,
    missing_findings: list[dict] | None,
    feedback: dict | None,
) -> dict:
    """Return the three independent work counts used by API/UI consumers."""
    rows = [item for item in (check_results or []) if isinstance(item, dict)]
    findings = [item for item in (missing_findings or []) if isinstance(item, dict)]
    problem_checks = sum(
        str(item.get("status") or "").lower() in {"fail", "warning", "error"}
        for item in rows
    )
    material_findings = sum(item.get("material_required", True) is not False for item in findings)
    return {
        "problem_checks": problem_checks,
        "material_findings": material_findings,
        "repair_tasks": int((feedback or {}).get("total", 0) or 0),
        "manual_or_later_findings": len(findings) - material_findings,
    }


def _chapter_natural_key(chapter_id: str) -> tuple:
    parts: list[int] = []
    for seg in str(chapter_id or "").split("."):
        try:
            parts.append(int(seg))
        except ValueError:
            parts.append(0)
    return tuple(parts)


def _fingerprint(chapter_id: str, detail: str) -> str:
    return f"{chapter_id}|{re.sub(r'[*_`\n]+', ' ', str(detail or ''))[:48]}"


def extract_from_check_results(results: dict | list | None) -> list[dict]:
    """从检查报告 results（{check_id: item} 或 [item]）提取 fact 型缺失 finding。"""
    if isinstance(results, dict):
        items = list(results.values())
    else:
        items = list(results or [])
    rows: list[dict] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        status = str(item.get("status", "")).lower()
        if status not in ("fail", "warning"):
            continue
        check_id = str(item.get("check_id", ""))
        check_name = str(item.get("check_name", ""))
        data = item.get("data") if isinstance(item.get("data"), dict) else {}
        findings = data.get("checks")
        if not isinstance(findings, list):
            findings = data.get("items") if isinstance(data.get("items"), list) else [data]
        candidates: list[dict] = []
        for f in findings:
            if isinstance(f, dict):
                candidates.append(f)
        candidates.append({"detail": item.get("reason", ""), "suggestion": ""})
        for f in candidates:
            if not isinstance(f, dict):
                continue
            f_status = str(f.get("status", status)).lower()
            if f_status not in ("fail", "warning"):
                continue
            detail = str(
                f.get("detail") or f.get("reason") or f.get("description") or f.get("title") or ""
            ).strip()
            if not detail or not is_missing_material_text(detail):
                continue
            # Warnings about an assumption, a coverage rate, or a cost model
            # are not upload tasks unless they name a concrete material/action.
            # This prevents “投标有效期未明确” and “成本说明不够细” from
            # appearing as enterprise evidence gaps.
            if check_id in {"validity_check", "pricing_logic_check", "fit_score"} and not re.search(
                r"材料|文件|证明|报告|证书|扫描件|附件|凭证|截图|上传|提供|签字|签章|盖章", detail
            ):
                continue
            detail = _sanitize_detail(detail)
            loc = _sanitize_detail(str(f.get("location") or f.get("chapter") or f.get("position") or "").strip())
            chapter_id = str(f.get("chapter_id") or "").strip()
            suggestion = _sanitize_detail(
                str(f.get("suggestion") or f.get("recommendation") or "").strip()
            )
            merged = f"{detail} {loc} {suggestion}"
            rows.append({
                "position": loc or (f"章节 {chapter_id}" if chapter_id else check_name or check_id),
                "chapter_id": chapter_id,
                "detail": detail[:300],
                "missing_type": classify_missing_type(merged, check_id, check_name),
                "suggested_material": suggestion or suggest_material(merged),
                "material_required": not is_non_current_material_matter(merged),
                "priority": priority_of(merged),
                "source": check_name or check_id or "检查报告",
            })
    return rows


def _sanitize_detail(text: str) -> str:
    """Worker J（净化层）：清单「发现详情」列不得携带生成期内部痕迹——
    【n】锚点/拒收原因（保留提取值）/知识库口径，统一过导出净化。
    章节正文 DB 原文不动；仅清单导出物净化。"""
    from core.agent_engine.export_sanitizer import sanitize_export_text

    cleaned, _ = sanitize_export_text(text)
    return cleaned


def extract_from_chapter_content(chapters: list[tuple]) -> list[dict]:
    """从章节正文提取【待补充】/（知识库无据）占位。chapters: [(id, title, content)]。"""
    rows: list[dict] = []
    for cid, title, content in chapters:
        text = str(content or "")
        if not text:
            continue
        for m in PEND_MARKER_RE.finditer(text):
            reason = (m.group(1) or "").strip("；; 　")
            start = max(0, m.start() - 50)
            ctx = re.sub(r"\s+", " ", text[start : m.start()]).strip()
            detail = f"…{ctx}【待补充】" if ctx else "【待补充】"
            if reason:
                detail = f"{detail}（{reason[:120]}）"
            detail = _sanitize_detail(detail)
            merged = f"{detail} {title}"
            rows.append({
                "position": f"{cid} {title}",
                "chapter_id": str(cid),
                "detail": detail[:300],
                "missing_type": classify_missing_type(merged),
                "suggested_material": suggest_material(merged),
                "material_required": not is_non_current_material_matter(merged),
                "priority": priority_of(merged),
                "source": "章节正文【待补充】标注",
            })
        for m in KB_MISSING_SENTENCE_RE.finditer(text):
            sentence = re.sub(r"\s+", " ", m.group(0)).strip()
            sentence = _sanitize_detail(sentence)
            merged = f"{sentence} {title}"
            rows.append({
                "position": f"{cid} {title}",
                "chapter_id": str(cid),
                "detail": sentence[:300],
                "missing_type": classify_missing_type(merged),
                "suggested_material": suggest_material(merged),
                "priority": priority_of(merged),
                "source": "章节正文无据占位",
            })
    return rows


def build_missing_materials_items(
    check_results: dict | list | None,
    chapters: list[tuple],
) -> list[dict]:
    """两路提取 + 去重（章节正文条目优先，检查报告补充）+ 排序（优先级→章节序）。"""
    from_chapters = extract_from_chapter_content(chapters)
    from_checks = extract_from_check_results(check_results)
    seen: set[str] = set()
    rows: list[dict] = []
    for row in from_chapters + from_checks:
        # “本阶段无需提交/后续阶段办理”是流程动作，不是企业材料缺口。
        # 在清单导出层再次过滤，防止旧报告或外部调用方遗漏 material_required 标记。
        if row.get("material_required") is False:
            continue
        key = _fingerprint(row.get("chapter_id", ""), row.get("detail", ""))
        if key in seen:
            continue
        seen.add(key)
        rows.append(row)
    rows.sort(key=lambda r: (_PRIORITY_ORDER.get(r["priority"], 2), _chapter_natural_key(r.get("chapter_id", ""))))
    for i, r in enumerate(rows, start=1):
        r["seq"] = i
    return rows


_PRIORITY_LABELS = {"high": "高", "medium": "中", "low": "低"}


def render_missing_materials_markdown(project_name: str, items: list[dict]) -> str:
    lines = [
        f"# {project_name}——需补充材料清单",
        "",
        f"共 {len(items)} 项待补充。按优先级排序；"
        "补齐材料后上传企业私有知识库，并回到检查页点击「资料已补充，重新检查修复」。",
        "",
        "| 序号 | 章节/位置 | 缺失要件类型 | 发现详情 | 建议收集的材料 | 优先级 | 来源 |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for r in items:
        detail = str(r.get("detail", "")).replace("|", "／").replace("\n", " ")
        material = str(r.get("suggested_material", "")).replace("|", "／")
        lines.append(
            f"| {r.get('seq', '')} | {r.get('position', '')} | {r.get('missing_type', '')} "
            f"| {detail} | {material} | {_PRIORITY_LABELS.get(r.get('priority', 'low'), '低')} "
            f"| {r.get('source', '')} |"
        )
    if not items:
        lines.append("| - | 无 | - | 未发现待补充材料项 | - | - | - |")
    return "\n".join(lines) + "\n"


def build_missing_materials_docx(project_name: str, items: list[dict]) -> bytes:
    """python-docx 表格版《需补充材料清单》；列：序号/章节位置/要件类型/发现详情/建议材料/优先级/来源。"""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(f"{project_name}——需补充材料清单", level=0)
    high = sum(1 for r in items if r.get("priority") == "high")
    medium = sum(1 for r in items if r.get("priority") == "medium")
    doc.add_paragraph(
        f"共 {len(items)} 项待补充（高优先级 {high} 项、中优先级 {medium} 项）。"
        "本清单由全量检查结果与章节【待补充】标注确定性生成（无 LLM）。"
        "请按行收集材料（签字/盖章/证书扫描件/最终报价等），上传企业私有知识库后，"
        "回到检查页点击「资料已补充，重新检查修复」完成闭环。"
    )
    headers = ["序号", "章节/位置", "缺失要件类型", "发现详情", "建议收集的材料", "优先级", "来源"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, head in zip(table.rows[0].cells, headers, strict=True):
        cell.text = head
    for r in items:
        cells = table.add_row().cells
        cells[0].text = str(r.get("seq", ""))
        cells[1].text = str(r.get("position", ""))
        cells[2].text = str(r.get("missing_type", ""))
        cells[3].text = str(r.get("detail", ""))
        cells[4].text = str(r.get("suggested_material", ""))
        cells[5].text = _PRIORITY_LABELS.get(str(r.get("priority", "low")), "低")
        cells[6].text = str(r.get("source", ""))
    if not items:
        cells = table.add_row().cells
        cells[1].text = "无待补充材料项"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = [
    "MISSING_RE",
    "is_missing_material_text",
    "classify_missing_type",
    "suggest_material",
    "priority_of",
    "reconcile_missing_findings",
    "build_action_summary",
    "extract_from_check_results",
    "extract_from_chapter_content",
    "build_missing_materials_items",
    "render_missing_materials_markdown",
    "build_missing_materials_docx",
]
