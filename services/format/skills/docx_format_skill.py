from __future__ import annotations

import logging
import os
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from core.skill_engine.base import Skill, SkillContext, SkillResult

logger = logging.getLogger(__name__)

RE_HAS_CHINESE = re.compile(r"[\u4e00-\u9fff]")
RE_MD_IMAGE = re.compile(r"!\[([^\]]*)\]\([^)]+\)")
RE_MD_LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
RE_MD_HTML_TAG = re.compile(r"<[^>]+>")
RE_MD_INLINE_CODE = re.compile(r"`([^`]+)`")
RE_MD_BOLD_ASTERISK = re.compile(r"\*\*(.*?)\*\*")
RE_MD_BOLD_UNDERSCORE = re.compile(r"__(.*?)__")
RE_MD_HEADER = re.compile(r"^\s*#+\s+(.*)")
RE_MD_BLOCKQUOTE = re.compile(r"^\s*>\s?(.*)")
RE_MD_HORIZONTAL_RULE = re.compile(r"^\s*[-*_]{3,}\s*$")
RE_MD_ORDERED_LIST = re.compile(r"^(\s*)(\d+)\.(\s+.*)")
RE_MD_UNORDERED_LIST_BLOCK = re.compile(r"^\s*[*+-]\s+")
RE_MD_UNORDERED_LIST = re.compile(r"^\s*[*+-]\s")
RE_MD_BULLET_WITH_CONTENT = re.compile(r"^(\s*[*+-]\s)(.*)")
RE_MD_EMPHASIS_ASTERISK = re.compile(r"(?<!\\)\*([^\s*][^*]*?)(?<!\\)\*")
RE_CURRENCY_PREFIX = re.compile(r"^[¥￥$]")
RE_CURRENCY_SUFFIX = re.compile(r"(元|万元|亿元)$")
RE_NUMERIC_TABLE_TEXT = re.compile(r"^[-+]?(?:\d+(?:\.\d+)?|\.\d+)%?$")
CHINESE_NUM_PATTERN = r"[一二三四五六七八九十百千万零]+"
RE_HEADING_H1 = re.compile(r"^" + CHINESE_NUM_PATTERN + r"\s*、")
RE_HEADING_H2 = re.compile(r"^[（\(]" + CHINESE_NUM_PATTERN + r"[）\)]")
RE_HEADING_H3 = re.compile(r"^\d+\s*[\.．]")
RE_HEADING_H4 = re.compile(r"^[（\(]\d+[）\)]")
RE_ATTACHMENT = re.compile(r"^附件\s*(\d+|[一二三四五六七八九十百千万零]+)?\s*[:：]?$")
RE_DECIMAL_H2 = re.compile(r"^(\d+)\s*[\.．]\s*(\d+)\s*[^\d\.．]")
RE_DECIMAL_H3 = re.compile(r"^(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[^\d\.．]")
RE_DECIMAL_H4 = re.compile(r"^(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[^\d\.．]")
RE_DECIMAL_H5 = re.compile(r"^(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[^\d\.．]")
RE_DECIMAL_H6 = re.compile(
    r"^(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[^\d\.．]"
)
RE_DECIMAL_H7 = re.compile(
    r"^(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[^\d\.．]"
)
RE_DECIMAL_H8 = re.compile(
    r"^(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[\.．]\s*(\d+)\s*[^\d\.．]"
)
RE_DECIMAL_HEADING_LIST = [
    RE_DECIMAL_H8,
    RE_DECIMAL_H7,
    RE_DECIMAL_H6,
    RE_DECIMAL_H5,
    RE_DECIMAL_H4,
    RE_DECIMAL_H3,
    RE_DECIMAL_H2,
]

STYLE_HIK_H4 = "15"
STYLE_HIK_H5 = "16"
STYLE_HIK_BODY = "11"
STYLE_HIK_INDENT = "25"
DEFAULT_INDENT_H4 = 360
DEFAULT_INDENT_H5 = 480

DEFAULT_FORMAT_CONFIG = {
    "page_number_align": "居中",
    "footer_distance": 2.5,
    "line_spacing": 28,
    "line_spacing_mode": "固定值(磅)",
    "margin_top": 2.6,
    "margin_bottom": 2.2,
    "margin_left": 2.5,
    "margin_right": 2.5,
    "title_font": "方正小标宋简体",
    "h1_font": "黑体",
    "h2_font": "黑体",
    "h3_font": "黑体",
    "h4_font": "楷体_GB2312",
    "h5_font": "楷体_GB2312",
    "h6_font": "宋体",
    "h7_font": "宋体",
    "h8_font": "宋体",
    "body_font": "仿宋_GB2312",
    "page_number_font": "宋体",
    "table_caption_font": "宋体",
    "figure_caption_font": "宋体",
    "attachment_font": "黑体",
    "subtitle_font": "楷体_GB2312",
    "table_font": "仿宋_GB2312",
    "table_header_font": "仿宋_GB2312",
    "english_font": "Times New Roman",
    "title_size": 22,
    "h1_size": 16,
    "h2_size": 15,
    "h3_size": 15,
    "h4_size": 14,
    "h5_size": 12,
    "h6_size": 10.5,
    "h7_size": 10.5,
    "h8_size": 10.5,
    "body_size": 12,
    "page_number_size": 14,
    "table_caption_size": 10.5,
    "figure_caption_size": 10.5,
    "attachment_size": 16,
    "subtitle_size": 16,
    "table_size": 12,
    "title_line_spacing": 33,
    "subtitle_line_spacing": 33,
    "h1_bold": True,
    "h2_bold": True,
    "h3_bold": True,
    "h4_bold": True,
    "h5_bold": True,
    "h6_bold": True,
    "h7_bold": False,
    "h8_bold": False,
    "h1_space_before": 12,
    "h1_space_after": 12,
    "h2_space_before": 6,
    "h2_space_after": 6,
    "h3_space_before": 6,
    "h3_space_after": 6,
    "h4_space_before": 2.4,
    "h4_space_after": 2.4,
    "h5_space_before": 1.2,
    "h5_space_after": 1.2,
    "left_indent_cm": 0.0,
    "right_indent_cm": 0.0,
    "set_outline": True,
    "enable_attachment_formatting": True,
    "force_a4": True,
    "use_custom_english_font": True,
    "normalize_punctuation": True,
    "enable_table_formatting": True,
    "table_auto_col_width": True,
    "table_width_percent": 100,
    "table_header_bold": True,
    "table_smart_align": True,
    "table_unified_borders": True,
    "table_border_size_pt": 0.5,
    "table_row_height_cm": 0.7,
    "table_line_spacing": 22,
    "heading_number_format": "decimal",
    "bold_before_colon": False,
    "indent_h4": DEFAULT_INDENT_H4,
    "indent_h5": DEFAULT_INDENT_H5,
}


class DocxFormatSkill(Skill):
    name = "docx_format"
    description = "智能排版(集成Word-Formatter-Pro引擎)"
    category = "output"
    version = "2.0.0"
    triggers = ["排版", "格式化", "格式调整", "美化", "格式检查"]

    async def execute(self, ctx: SkillContext) -> SkillResult:
        file_path = ctx.parameters.get("file_path", "")

        mode = ctx.parameters.get("mode", "format")
        custom_config = ctx.parameters.get("config", {})

        if not file_path or not Path(file_path).exists():
            return SkillResult(success=False, error=f"文件不存在: {file_path}")

        config = {**DEFAULT_FORMAT_CONFIG, **custom_config}

        if mode == "check":
            return await self._check_format(file_path, config)
        elif mode == "diff":
            return await self._diff_format(file_path, config)
        elif mode == "beautify":
            return await self._beautify_xml(file_path, config)
        else:
            return await self._format_document(file_path, config)

    async def _check_format(self, file_path: str, config: dict) -> SkillResult:
        from docx import Document

        try:
            doc = Document(file_path)
        except Exception as e:
            return SkillResult(success=False, error=f"文件打开失败: {e}")

        issues: list[dict[str, Any]] = []
        total_paragraphs = 0
        checked_paragraphs = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            total_paragraphs += 1

            expected_label = self._detect_heading_label(text)
            expected_rules = self._get_rules_for_label(expected_label, config)

            for run in para.runs:
                checked_paragraphs += 1
                if "font_name" in expected_rules:
                    actual_font = run.font.name or ""
                    east_asia_font = ""
                    try:
                        from docx.oxml.ns import qn

                        r_pr = run._element.r_pr
                        if r_pr is not None:
                            r_fonts = r_pr.find(qn("w:rFonts"))
                            if r_fonts is not None:
                                east_asia_font = r_fonts.get(qn("w:eastAsia"), "")
                    except Exception:
                        pass
                    expected = expected_rules["font_name"]
                    if actual_font != expected and east_asia_font != expected:
                        issues.append(
                            {
                                "type": "font_mismatch",
                                "expected": expected,
                                "actual": actual_font or east_asia_font,
                                "text_preview": text[:30],
                                "label": expected_label,
                            }
                        )

                if "font_size" in expected_rules:
                    actual_size = run.font.size
                    if actual_size and actual_size.pt != expected_rules["font_size"]:
                        issues.append(
                            {
                                "type": "size_mismatch",
                                "expected": expected_rules["font_size"],
                                "actual": actual_size.pt,
                                "text_preview": text[:30],
                                "label": expected_label,
                            }
                        )

                if "bold" in expected_rules:
                    if run.font.bold != expected_rules["bold"]:
                        issues.append(
                            {
                                "type": "bold_mismatch",
                                "expected": expected_rules["bold"],
                                "actual": run.font.bold,
                                "text_preview": text[:30],
                                "label": expected_label,
                            }
                        )

            if expected_label == "正文" and "first_line_indent" in expected_rules:
                pf = para.paragraph_format
                if pf.first_line_indent is None or pf.first_line_indent.pt == 0:
                    issues.append(
                        {
                            "type": "indent_missing",
                            "expected": "首行缩进2字符",
                            "actual": "无缩进",
                            "text_preview": text[:30],
                            "label": expected_label,
                        }
                    )

        for section in doc.sections:
            if abs(section.top_margin.cm - config["margin_top"]) > 0.1:
                issues.append(
                    {
                        "type": "margin_mismatch",
                        "expected": f"{config['margin_top']}cm",
                        "actual": f"{section.top_margin.cm:.1f}cm",
                        "detail": "上边距",
                    }
                )
            if abs(section.bottom_margin.cm - config["margin_bottom"]) > 0.1:
                issues.append(
                    {
                        "type": "margin_mismatch",
                        "expected": f"{config['margin_bottom']}cm",
                        "actual": f"{section.bottom_margin.cm:.1f}cm",
                        "detail": "下边距",
                    }
                )
            if abs(section.left_margin.cm - config["margin_left"]) > 0.1:
                issues.append(
                    {
                        "type": "margin_mismatch",
                        "expected": f"{config['margin_left']}cm",
                        "actual": f"{section.left_margin.cm:.1f}cm",
                        "detail": "左边距",
                    }
                )
            if abs(section.right_margin.cm - config["margin_right"]) > 0.1:
                issues.append(
                    {
                        "type": "margin_mismatch",
                        "expected": f"{config['margin_right']}cm",
                        "actual": f"{section.right_margin.cm:.1f}cm",
                        "detail": "右边距",
                    }
                )

        font_issues = [i for i in issues if i["type"] == "font_mismatch"]
        size_issues = [i for i in issues if i["type"] == "size_mismatch"]
        indent_issues = [i for i in issues if i["type"] == "indent_missing"]
        margin_issues = [i for i in issues if i["type"] == "margin_mismatch"]

        return SkillResult(
            success=True,
            data={
                "total_paragraphs": total_paragraphs,
                "checked_runs": checked_paragraphs,
                "total_issues": len(issues),
                "font_issues": len(font_issues),
                "size_issues": len(size_issues),
                "indent_issues": len(indent_issues),
                "margin_issues": len(margin_issues),
                "issues": issues[:100],
                "compliance_rate": max(0, round((1 - len(issues) / max(checked_paragraphs, 1)) * 100, 1)),
            },
        )

    async def _diff_format(self, file_path: str, config: dict) -> SkillResult:
        check_result = await self._check_format(file_path, config)
        if not check_result.success:
            return check_result

        diff_items = []
        for issue in check_result.data.get("issues", []):
            diff_items.append(
                {
                    "location": issue.get("text_preview", issue.get("detail", "")),
                    "property": issue.get("type", ""),
                    "expected": issue.get("expected", ""),
                    "actual": issue.get("actual", ""),
                    "label": issue.get("label", ""),
                }
            )

        return SkillResult(
            success=True,
            data={
                "diff_count": len(diff_items),
                "diffs": diff_items[:200],
                "auto_fixable": len(
                    [
                        d
                        for d in diff_items
                        if d["property"]
                        in ("font_mismatch", "size_mismatch", "bold_mismatch", "indent_missing", "margin_mismatch")
                    ]
                ),
                "check_summary": {
                    "total_paragraphs": check_result.data["total_paragraphs"],
                    "total_issues": check_result.data["total_issues"],
                    "compliance_rate": check_result.data["compliance_rate"],
                },
            },
        )

    async def _format_document(self, file_path: str, config: dict) -> SkillResult:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm

        try:
            doc = Document(file_path)
        except Exception as e:
            return SkillResult(success=False, error=f"文件打开失败: {e}")

        stats = {"paragraphs": 0, "headings": 0, "body": 0, "tables": 0, "margins_fixed": 0}

        for section in doc.sections:
            section.top_margin = Cm(config["margin_top"])
            section.bottom_margin = Cm(config["margin_bottom"])
            section.left_margin = Cm(config["margin_left"])
            section.right_margin = Cm(config["margin_right"])
            section.footer_distance = Cm(config["footer_distance"])
            if config.get("force_a4", False):
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
            stats["margins_fixed"] += 1

            if config.get("page_number_align") == "居中":
                footer = section.footer
                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._create_page_number(p, config)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            stats["paragraphs"] += 1

            label = self._detect_heading_label(text)
            rules = self._get_rules_for_label(label, config)

            self._apply_format_to_para(para, rules, config)

            if label != "正文":
                stats["headings"] += 1
            else:
                stats["body"] += 1

        # Worker J（净化层）：排版交付物统一导出净化——【n】锚点/硬门拒收原因/
        # 插图建议脚手架不得进入最终 DOCX（数据库原文不动，只净化导出物）。
        try:
            from core.agent_engine.export_sanitizer import sanitize_docx_paragraphs

            stats["sanitized"] = sanitize_docx_paragraphs(doc)
        except Exception as e:  # noqa: BLE001
            logger.warning("排版 DOCX 导出净化失败（降级原文输出）: %s", e)
            stats["sanitized"] = {}

        if config.get("enable_table_formatting", True):
            for table in doc.tables:
                stats["tables"] += 1
                self._format_table(table, config)

        output_path = str(Path(file_path).with_suffix("")) + "_formatted.docx"
        doc.save(output_path)

        # G-7 收尾：配图增量——正文中 Markdown 图片引用（含 data URI）原位嵌图。
        # 开关关/无图引用时为 0，行为不变；开关开 + 生成期装配过图片引用时导出即嵌图。
        try:
            from core.settings import get_settings

            if get_settings().illustration_enabled:
                from core.agent_engine.illustration_assembly import embed_images_into_docx

                stats["illustrations_embedded"] = embed_images_into_docx(output_path)
        except Exception as e:  # noqa: BLE001
            stats["illustrations_embedded"] = 0
            stats["illustrations_error"] = str(e)[:200]

        return SkillResult(
            success=True,
            data={
                "output_path": output_path,
                "stats": stats,
            },
        )

    async def _beautify_xml(self, file_path: str, config: dict) -> SkillResult:
        tmp = tempfile.mkdtemp(prefix="bidfmt_")
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                zf.extractall(tmp)

            doc_path = os.path.join(tmp, "word", "document.xml")
            num_path = os.path.join(tmp, "word", "numbering.xml")
            sty_path = os.path.join(tmp, "word", "styles.xml")

            report: dict[str, Any] = {}

            if os.path.exists(doc_path):
                with open(doc_path, "r", encoding="utf-8") as f:
                    doc_xml = f.read()

                doc_xml, h4_cnt, h5_cnt = self._step_apply_heading_styles(doc_xml)
                doc_xml, dup_stripped = self._step_strip_duplicate_prefixes(doc_xml)
                doc_xml, body_cnt = self._step_apply_body_indent(doc_xml)

                with open(doc_path, "w", encoding="utf-8") as f:
                    f.write(doc_xml)

                report["h4_applied"] = h4_cnt
                report["h5_applied"] = h5_cnt
                report["dup_stripped"] = dup_stripped
                report["body_indent"] = body_cnt

            if os.path.exists(num_path):
                with open(num_path, "r", encoding="utf-8") as f:
                    num_xml = f.read()
                indent_h4 = config.get("indent_h4", DEFAULT_INDENT_H4)
                indent_h5 = config.get("indent_h5", DEFAULT_INDENT_H5)
                num_xml, num_changes = self._optimize_numbering(num_xml, indent_h4, indent_h5)
                with open(num_path, "w", encoding="utf-8") as f:
                    f.write(num_xml)
                report["num_optimizations"] = len(num_changes)

            if os.path.exists(sty_path):
                with open(sty_path, "r", encoding="utf-8") as f:
                    sty_xml = f.read()
                sty_xml, sty_changes = self._optimize_styles(sty_xml)
                with open(sty_path, "w", encoding="utf-8") as f:
                    f.write(sty_xml)
                report["style_optimizations"] = len(sty_changes)

            output_path = str(Path(file_path).with_suffix("")) + "_beautified.docx"
            with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(tmp):
                    for fname in files:
                        fpath = os.path.join(root, fname)
                        arcname = os.path.relpath(fpath, tmp)
                        zf.write(fpath, arcname)

            return SkillResult(
                success=True,
                data={"output_path": output_path, "report": report},
            )
        except Exception as e:
            return SkillResult(success=False, error=f"XML美化失败: {e}")
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    def _detect_heading_label(self, text: str) -> str:
        heading_format = DEFAULT_FORMAT_CONFIG.get("heading_number_format", "decimal")
        if heading_format == "decimal":
            for i, regex in enumerate(RE_DECIMAL_HEADING_LIST):
                if regex.match(text):
                    level = 8 - i
                    return f"h{level}"
        elif heading_format == "decimal_paren":
            if re.match(r"^\d+\s*[）\)]", text):
                return "h1"
            if re.match(r"^\d+\s*[\.．]\s*\d+\s*[）\)]", text):
                return "h2"
            if re.match(r"^\d+\s*[\.．]\s*\d+\s*[\.．]\s*\d+\s*[）\)]", text):
                return "h3"
        elif heading_format == "chinese_paren":
            if re.match(r"^[（\(]" + CHINESE_NUM_PATTERN + r"\s*[）\)]", text):
                return "h2"
        elif heading_format == "mixed":
            if RE_HEADING_H1.match(text):
                return "h1"
            for i, regex in enumerate(RE_DECIMAL_HEADING_LIST):
                if regex.match(text):
                    level = 8 - i
                    return f"h{level}"
        elif heading_format == "letter":
            if re.match(r"^[A-Z]\s*[\.．、]", text):
                return "h1"
            if re.match(r"^[a-z]\s*[\.．、]", text):
                return "h2"
            if re.match(r"^[A-Z]\s*[\.．]\s*\d+", text):
                return "h2"
            if re.match(r"^[A-Z]\s*[\.．]\s*\d+\s*[\.．]\s*\d+", text):
                return "h3"
        if RE_HEADING_H1.match(text):
            return "h1"
        if RE_HEADING_H2.match(text):
            return "h2"
        if RE_HEADING_H3.match(text):
            return "h3"
        if RE_HEADING_H4.match(text):
            return "h4"
        if RE_ATTACHMENT.match(text):
            return "attachment"
        return "正文"

    def _get_rules_for_label(self, label: str, config: dict) -> dict:
        if label.startswith("h"):
            level = label[1:]
            return {
                "font_name": config.get(f"h{level}_font", config.get("body_font", "宋体")),
                "font_size": config.get(f"h{level}_size", 12),
                "bold": config.get(f"h{level}_bold", True),
                "space_before": config.get(f"h{level}_space_before", 0),
                "space_after": config.get(f"h{level}_space_after", 0),
                "alignment": "left",
                "first_line_indent": 0,
            }
        if label == "attachment":
            return {
                "font_name": config.get("attachment_font", "黑体"),
                "font_size": config.get("attachment_size", 16),
                "bold": True,
                "alignment": "left",
                "first_line_indent": 0,
            }
        return {
            "font_name": config.get("body_font", "仿宋_GB2312"),
            "font_size": config.get("body_size", 12),
            "bold": False,
            "alignment": "justify",
            "first_line_indent": 2,
            "line_spacing": config.get("line_spacing", 28),
        }

    def _apply_format_to_para(self, para, rules: dict, config: dict):
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        if not para.runs:
            return

        en_font = (
            config.get("english_font", "Times New Roman")
            if config.get("use_custom_english_font", False)
            else rules.get("font_name", "宋体")
        )

        for run in para.runs:
            if "font_name" in rules:
                run.font.name = en_font or rules["font_name"]
                r_pr = run._element.get_or_add_rPr()
                r_fonts = r_pr.find(qn("w:rFonts"))
                if r_fonts is None:
                    r_fonts = OxmlElement("w:rFonts")
                    r_pr.insert(0, r_fonts)
                for theme_attr in ("w:eastAsiaTheme", "w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:csTheme"):
                    r_fonts.attrib.pop(qn(theme_attr), None)
                r_fonts.set(qn("w:eastAsia"), rules["font_name"])
                r_fonts.set(qn("w:ascii"), en_font or rules["font_name"])
                r_fonts.set(qn("w:hAnsi"), en_font or rules["font_name"])
            if "font_size" in rules:
                run.font.size = Pt(rules["font_size"])
            if "bold" in rules:
                run.font.bold = rules["bold"]

        if "alignment" in rules:
            align_map = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(rules["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

        if "first_line_indent" in rules:
            if rules["first_line_indent"] == 0:
                para.paragraph_format.first_line_indent = None
                p_pr = para._p.get_or_add_pPr()
                ind = p_pr.find(qn("w:ind"))
                if ind is not None:
                    ind.set(qn("w:firstLineChars"), "0")
                    for attr in (qn("w:firstLine"), qn("w:hanging"), qn("w:hangingChars")):
                        ind.attrib.pop(attr, None)
            elif "font_size" in rules:
                para.paragraph_format.first_line_indent = Pt(rules["font_size"] * rules["first_line_indent"] * 2)

        if "space_before" in rules:
            para.paragraph_format.space_before = Pt(rules["space_before"])
        if "space_after" in rules:
            para.paragraph_format.space_after = Pt(rules["space_after"])
        if "line_spacing" in rules:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(rules["line_spacing"])

        if config.get("set_outline", False) and rules.get("alignment") == "left":
            label = self._detect_heading_label(para.text.strip())
            if label.startswith("h"):
                level = int(label[1:])
                p_pr = para._p.get_or_add_pPr()
                outline_lvl = p_pr.find(qn("w:outlineLvl"))
                if outline_lvl is None:
                    outline_lvl = OxmlElement("w:outlineLvl")
                    p_pr.append(outline_lvl)
                outline_lvl.set(qn("w:val"), str(level - 1))

    def _create_page_number(self, paragraph, config: dict):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        font_name = config.get("page_number_font", "宋体")
        font_size = config.get("page_number_size", 14)

        run_prefix = paragraph.add_run("— ")
        run_prefix.font.size = Pt(font_size)
        run_prefix.font.color.rgb = RGBColor(0, 0, 0)
        run_prefix.font.name = font_name

        run_field = paragraph.add_run()
        run_field.font.size = Pt(font_size)
        run_field.font.color.rgb = RGBColor(0, 0, 0)
        run_field.font.name = font_name

        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run_field._r.extend([fld_char1, instr_text, fld_char2])

        run_suffix = paragraph.add_run(" —")
        run_suffix.font.size = Pt(font_size)
        run_suffix.font.color.rgb = RGBColor(0, 0, 0)
        run_suffix.font.name = font_name

    def _format_table(self, table, config: dict):
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        table_font = config.get("table_font", "仿宋_GB2312")
        table_header_font = config.get("table_header_font", table_font)
        table_size = config.get("table_size", 12)
        table_line_spacing = config.get("table_line_spacing", 22)
        row_height_cm = config.get("table_row_height_cm", 0.7)
        border_size_pt = config.get("table_border_size_pt", 0.5)
        header_bold = config.get("table_header_bold", True)
        smart_align = config.get("table_smart_align", False)

        size = max(1, int(float(border_size_pt) * 8))
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        else:
            for child in list(borders):
                borders.remove(child)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            elem = OxmlElement(f"w:{edge}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), str(size))
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "000000")
            borders.append(elem)

        for row_idx, row in enumerate(table.rows):
            if row_height_cm > 0:
                row.height = Cm(row_height_cm)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        for run in para.runs:
                            font_name = table_header_font if row_idx == 0 else table_font
                            run.font.name = font_name
                            r_pr = run._element.get_or_add_rPr()
                            r_fonts = r_pr.find(qn("w:rFonts"))
                            if r_fonts is None:
                                r_fonts = OxmlElement("w:rFonts")
                                r_pr.insert(0, r_fonts)
                            r_fonts.set(qn("w:eastAsia"), font_name)
                            r_fonts.set(qn("w:ascii"), font_name)
                            r_fonts.set(qn("w:hAnsi"), font_name)
                            run.font.size = Pt(table_size)
                            if row_idx == 0 and header_bold:
                                run.font.bold = True
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    if table_line_spacing > 0:
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        para.paragraph_format.line_spacing = Pt(table_line_spacing)
                    if smart_align:
                        if row_idx == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif self._is_numeric_text(para.text.strip()):
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def _is_numeric_text(text: str) -> bool:
        text = text.strip().replace(",", "").replace("，", "").replace("％", "%")
        text = RE_CURRENCY_PREFIX.sub("", text)
        text = RE_CURRENCY_SUFFIX.sub("", text)
        return bool(RE_NUMERIC_TABLE_TEXT.match(text))

    def _step_apply_heading_styles(self, content: str) -> tuple[str, int, int]:
        para_pattern = r"(<w:p[ >])(.*?)(</w:p>)"
        paras = list(re.finditer(para_pattern, content, re.DOTALL))
        counters = {"h4": 0, "h5": 0}
        modifications = []

        for m in paras:
            p_body = m.group(2)
            if re.search(r"<w:pStyle", p_body):
                continue
            full_text = self._get_text_from_para(p_body)
            if not full_text:
                continue

            m5 = re.match(r"^(\d+)\.(\d+)\.(\d+)\.(\d+)\.(\d+)[ \t]*(.*)$", full_text)
            if m5 and self._is_likely_heading(m5.group(6)):
                new_p_pr = self._insert_pstyle(p_body, STYLE_HIK_H5)
                if new_p_pr != p_body:
                    counters["h5"] += 1
                    modifications.append((m.start(2), m.end(2), new_p_pr))
                continue

            m4 = re.match(r"^(\d+)\.(\d+)\.(\d+)\.(\d+)([ \t]*)(.*)$", full_text)
            if m4 and self._is_likely_heading(m4.group(6)):
                new_p_pr = self._insert_pstyle(p_body, STYLE_HIK_H4)
                if new_p_pr != p_body:
                    counters["h4"] += 1
                    modifications.append((m.start(2), m.end(2), new_p_pr))

        content_list = list(content)
        for start, end, repl in reversed(modifications):
            content_list[start:end] = list(repl)
        return "".join(content_list), counters["h4"], counters["h5"]

    def _step_strip_duplicate_prefixes(self, content: str) -> tuple[str, int]:
        para_pattern = r"(<w:p[ >])(.*?)(</w:p>)"
        paras = list(re.finditer(para_pattern, content, re.DOTALL))
        stripped = 0
        modifications = []

        for m in paras:
            p_body = m.group(2)
            if not re.search(rf'w:val="({STYLE_HIK_H4}|{STYLE_HIK_H5})"', p_body):
                continue
            t_match = re.search(r"(<w:t[^>]*>)([^<]*)(</w:t>)", p_body)
            if not t_match:
                continue
            text_val = t_match.group(2).strip()
            if not text_val:
                continue
            new_text = re.sub(r"^\d+(\.\d+){3,}[ \t]*(?=.)", "", text_val, count=1)
            if new_text != text_val and new_text.strip():
                new_full = t_match.group(1) + new_text + t_match.group(3)
                modifications.append((m.start(2) + t_match.start(), m.start(2) + t_match.end(), new_full))
                stripped += 1

        content_list = list(content)
        for start, end, repl in reversed(modifications):
            content_list[start:end] = list(repl)
        return "".join(content_list), stripped

    def _step_apply_body_indent(self, content: str) -> tuple[str, int]:
        para_pattern = r"(<w:p[ >])(.*?)(</w:p>)"
        paras = list(re.finditer(para_pattern, content, re.DOTALL))
        body_count = 0
        modifications = []

        for m in paras:
            p_body = m.group(2)
            if re.search(r"<w:pStyle", p_body):
                continue
            full_text = self._get_text_from_para(p_body)
            if not full_text:
                continue
            if re.match(r"^[\d\s.,;:\-+=\\/\\(){}\[\]\"\'<>]+$", full_text):
                continue
            if re.match(r"^\d+\.", full_text):
                continue
            new_p_pr = self._insert_pstyle(p_body, STYLE_HIK_INDENT)
            if new_p_pr != p_body:
                body_count += 1
                modifications.append((m.start(2), m.end(2), new_p_pr))

        content_list = list(content)
        for start, end, repl in reversed(modifications):
            content_list[start:end] = list(repl)
        return "".join(content_list), body_count

    def _optimize_numbering(self, num_xml: str, indent_h4: int, indent_h5: int) -> tuple[str, list[str]]:
        changes = []

        def replace_ind(xml, abs_id, ilvl, new_left, label):
            pat = (
                rf'(<w:abstractNum w:abstractNumId="{abs_id}"[^>]*>.*?'
                rf'<w:lvl w:ilvl="{ilvl}"[^>]*>.*?'
                r'<w:ind w:left=")(\d+)(" w:hanging=")(\d+)("/>)'
            )
            m = re.search(pat, xml, re.DOTALL)
            if m:
                old_l, old_h = int(m.group(2)), int(m.group(4))
                xml = re.sub(pat, rf"\g<1>{new_left}\g<3>{new_left}\g<5>", xml, count=1, flags=re.DOTALL)
                changes.append(f"{label}: left={old_l}→{new_left}, hanging={old_h}→{new_left}")
            return xml

        result = num_xml
        result = replace_ind(result, "1", "3", indent_h4, "H4缩进(ilvl=3)")
        result = replace_ind(result, "1", "4", indent_h5, "H5缩进(ilvl=4)")
        return result, changes

    def _optimize_styles(self, style_xml: str) -> tuple[str, list[str]]:
        changes = []
        result = style_xml

        old_s = r'<w:spacing w:before="120" w:after="120" w:line="360" w:lineRule="auto"/>'
        new_s = '<w:spacing w:before="60" w:after="40" w:line="280" w:lineRule="auto"/>'
        if re.search(old_s, result):
            result = re.sub(old_s, new_s, result)
            changes.append("hik标题4: 段前6pt→3pt, 段后6pt→2pt, 行距280")

        old_f = r'<w:rFonts w:ascii="黑体" w:hAnsi="宋体" w:eastAsia="黑体" w:cs="Times New Roman"/>'
        new_f = '<w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体" w:cs="黑体"/>'
        if re.search(old_f, result):
            result = re.sub(old_f, new_f, result)
            changes.append("hik标题4: 字体统一黑体")

        m5ppr = re.search(
            r'(<w:style w:type="paragraph"[^>]*w:styleId="16"[^>]*>.*?<w:pPr>)(.*?)(</w:pPr>)', result, re.DOTALL
        )
        if m5ppr:
            pc = m5ppr.group(2)
            if not re.search("<w:spacing", pc):
                npc = pc + '<w:spacing w:before="40" w:after="20" w:line="276" w:lineRule="auto"/>'
                result = result[: m5ppr.start(2)] + npc + result[m5ppr.end(2) :]
                changes.append("hik标题5: 新增段前2pt/段后1pt/行距")

        old_h5f = r'<w:rFonts w:ascii="宋体" w:hAnsi="黑体" w:eastAsia="宋体" w:cs="宋体"/>'
        new_h5f = '<w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体" w:cs="黑体"/>'
        if re.search(old_h5f, result):
            result = re.sub(old_h5f, new_h5f, result)
            changes.append("hik标题5: 字体改为黑体")

        m5rpr = re.search(
            r'(<w:style w:type="paragraph"[^>]*w:styleId="16"[^>]*>.*?<w:rPr>)(.*?)(</w:rPr>)', result, re.DOTALL
        )
        if m5rpr and not re.search("<w:sz ", m5rpr.group(2)):
            nr = m5rpr.group(2) + '<w:sz w:val="21"/><w:szCs w:val="21"/>'
            result = result[: m5rpr.start(2)] + nr + result[m5rpr.end(2) :]
            changes.append("hik标题5: 新增字号10.5pt")

        return result, changes

    @staticmethod
    def _get_text_from_para(para_body: str) -> str:
        texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", para_body)
        return "".join(texts).strip()

    @staticmethod
    def _is_likely_heading(body_text: str) -> bool:
        t = body_text.strip()
        if not t or len(t) > 60:
            return False
        if re.search(r"[。；！？]", t) and len(t) > 15:
            return False
        if re.match(r"^[a-zA-Z0-9\-\+,./\\()[\]]{2,}$", t) and len(t) < 20:
            return False
        return True

    @staticmethod
    def _insert_pstyle(p_body: str, style_id: str) -> str:
        if re.search(r"<w:pStyle\s", p_body):
            return p_body
        if "<w:pPr>" in p_body:
            return p_body.replace("<w:pPr>", f'<w:pPr><w:pStyle w:val="{style_id}"/>')
        return f'<w:pPr><w:pStyle w:val="{style_id}"/></w:pPr>{p_body}'
