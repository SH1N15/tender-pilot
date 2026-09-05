from __future__ import annotations

import logging
import os
import shutil
import subprocess
from pathlib import Path
from typing import Any

from core.skill_engine.base import Skill, SkillContext, SkillResult

logger = logging.getLogger(__name__)

WORD_HEADING_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "标题 1": 1,
    "标题 2": 2,
    "标题 3": 3,
    "标题 4": 4,
    "标题 5": 5,
    "标题 6": 6,
}


class PdfExportSkill(Skill):
    name = "pdf_export"
    description = "Word转PDF导出，支持书签/目录/超链接保留"
    category = "output"
    version = "1.0.0"
    triggers = ["PDF导出", "转PDF", "Word转PDF", "导出PDF"]

    async def execute(self, ctx: SkillContext) -> SkillResult:
        input_path = ctx.parameters.get("input_path", "")
        output_dir = ctx.parameters.get("output_dir", "")

        if not input_path:
            return SkillResult(success=False, error="未提供输入文件路径(input_path)")

        if not Path(input_path).exists():
            return SkillResult(success=False, error=f"输入文件不存在: {input_path}")

        if not input_path.lower().endswith(".docx"):
            return SkillResult(success=False, error="输入文件必须是.docx格式")

        bookmarks = self._extract_bookmarks(input_path)

        conversion_result = await self._convert_to_pdf(input_path, output_dir)
        if not conversion_result["success"]:
            return SkillResult(
                success=False,
                error=conversion_result["error"],
                warnings=conversion_result.get("warnings", []),
            )

        output_path = conversion_result["output_path"]
        method_used = conversion_result["method"]

        page_count = self._count_pdf_pages(output_path)

        bookmark_info = self._build_bookmark_info(bookmarks)

        return SkillResult(
            success=True,
            data={
                "output_path": output_path,
                "page_count": page_count,
                "bookmarks": bookmark_info,
                "bookmark_count": len(bookmarks),
                "method": method_used,
            },
            warnings=conversion_result.get("warnings", []),
        )

    def _extract_bookmarks(self, input_path: str) -> list[dict[str, Any]]:
        from docx import Document

        bookmarks: list[dict[str, Any]] = []
        try:
            doc = Document(input_path)
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else ""
                level = WORD_HEADING_LEVELS.get(style_name, 0)
                if level > 0:
                    text = para.text.strip()
                    if text:
                        bookmarks.append(
                            {
                                "title": text,
                                "level": level,
                            }
                        )
        except Exception as e:
            logger.warning("提取书签信息失败: %s", e)

        return bookmarks

    async def _convert_to_pdf(self, input_path: str, output_dir: str) -> dict[str, Any]:
        warnings: list[str] = []

        result = self._try_libreoffice(input_path, output_dir)
        if result["success"]:
            result["warnings"] = warnings
            return result
        warnings.append(f"LibreOffice转换失败: {result.get('error', '未知错误')}")

        result = self._try_docx2pdf(input_path, output_dir)
        if result["success"]:
            result["warnings"] = warnings
            return result
        warnings.append(f"docx2pdf转换失败: {result.get('error', '未知错误')}")

        result = await self._try_reportlab(input_path, output_dir)
        if result["success"]:
            result["warnings"] = warnings
            return result
        warnings.append(f"reportlab转换失败: {result.get('error', '未知错误')}")

        return {
            "success": False,
            "error": "所有PDF转换策略均失败(LibreOffice/docx2pdf/reportlab)",
            "warnings": warnings,
        }

    def _try_libreoffice(self, input_path: str, output_dir: str) -> dict[str, Any]:
        soffice_paths = [
            "soffice",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]

        soffice_cmd = None
        for candidate in soffice_paths:
            if shutil.which(candidate):
                soffice_cmd = candidate
                break
            if os.path.isfile(candidate):
                soffice_cmd = candidate
                break

        if not soffice_cmd:
            return {"success": False, "error": "未找到LibreOffice(soffice)"}

        work_dir = output_dir or str(Path(input_path).parent)

        try:
            cmd = [
                soffice_cmd,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                work_dir,
                input_path,
            ]
            proc = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
            )

            if proc.returncode != 0:
                return {
                    "success": False,
                    "error": f"LibreOffice返回非零退出码: {proc.returncode}, stderr: {proc.stderr[:500]}",
                }

            pdf_name = Path(input_path).stem + ".pdf"
            pdf_path = os.path.join(work_dir, pdf_name)

            if not Path(pdf_path).exists():
                return {"success": False, "error": "LibreOffice转换后未找到PDF文件"}

            return {
                "success": True,
                "output_path": pdf_path,
                "method": "libreoffice",
            }
        except subprocess.TimeoutExpired:
            return {"success": False, "error": "LibreOffice转换超时(120s)"}
        except Exception as e:
            return {"success": False, "error": f"LibreOffice异常: {e}"}

    def _try_docx2pdf(self, input_path: str, output_dir: str) -> dict[str, Any]:
        try:
            import docx2pdf
        except ImportError:
            return {"success": False, "error": "docx2pdf库未安装"}

        try:
            if output_dir:
                pdf_path = os.path.join(output_dir, Path(input_path).stem + ".pdf")
                docx2pdf.convert(input_path, pdf_path)
            else:
                docx2pdf.convert(input_path)
                pdf_path = str(Path(input_path).with_suffix(".pdf"))

            if not Path(pdf_path).exists():
                return {"success": False, "error": "docx2pdf转换后未找到PDF文件"}

            return {
                "success": True,
                "output_path": pdf_path,
                "method": "docx2pdf",
            }
        except Exception as e:
            return {"success": False, "error": f"docx2pdf异常: {e}"}

    async def _try_reportlab(self, input_path: str, output_dir: str) -> dict[str, Any]:
        try:
            from docx import Document
            from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as e:
            return {"success": False, "error": f"依赖库未安装: {e}"}

        try:
            doc = Document(input_path)
        except Exception as e:
            return {"success": False, "error": f"读取DOCX失败: {e}"}

        if output_dir:
            pdf_path = os.path.join(output_dir, Path(input_path).stem + ".pdf")
        else:
            pdf_path = str(Path(input_path).with_suffix(".pdf"))

        cn_font_name = self._register_chinese_font()

        try:
            pdf_doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                leftMargin=2.5 * cm,
                rightMargin=2.5 * cm,
                topMargin=2.6 * cm,
                bottomMargin=2.2 * cm,
            )

            styles_config = {
                "h1": ParagraphStyle(
                    "H1",
                    fontName=cn_font_name,
                    fontSize=16,
                    leading=24,
                    spaceBefore=12,
                    spaceAfter=6,
                    alignment=TA_LEFT,
                ),
                "h2": ParagraphStyle(
                    "H2",
                    fontName=cn_font_name,
                    fontSize=15,
                    leading=22,
                    spaceBefore=10,
                    spaceAfter=5,
                    alignment=TA_LEFT,
                ),
                "h3": ParagraphStyle(
                    "H3",
                    fontName=cn_font_name,
                    fontSize=14,
                    leading=20,
                    spaceBefore=8,
                    spaceAfter=4,
                    alignment=TA_LEFT,
                ),
                "h4": ParagraphStyle(
                    "H4",
                    fontName=cn_font_name,
                    fontSize=12,
                    leading=18,
                    spaceBefore=6,
                    spaceAfter=3,
                    alignment=TA_LEFT,
                ),
                "body": ParagraphStyle(
                    "Body",
                    fontName=cn_font_name,
                    fontSize=12,
                    leading=20,
                    spaceBefore=2,
                    spaceAfter=2,
                    firstLineIndent=24,
                    alignment=TA_JUSTIFY,
                ),
            }

            story: list[Any] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    story.append(Spacer(1, 6))
                    continue

                style_name = para.style.name if para.style else ""
                level = WORD_HEADING_LEVELS.get(style_name, 0)

                escaped_text = self._escape_xml(text)

                if level > 0:
                    style_key = f"h{min(level, 4)}"
                    style = styles_config[style_key]
                    if level <= 2:
                        story.append(Spacer(1, 8))
                    story.append(Paragraph(escaped_text, style))
                else:
                    story.append(Paragraph(escaped_text, styles_config["body"]))

            for table in doc.tables:
                story.append(Spacer(1, 8))
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip() or " "
                        row_data.append(
                            Paragraph(
                                self._escape_xml(cell_text),
                                styles_config["body"],
                            )
                        )
                    table_data.append(row_data)

                if table_data:
                    from reportlab.platypus import Table as RLTable

                    col_count = max(len(r) for r in table_data)
                    available_width = A4[0] - 5 * cm
                    col_width = available_width / max(col_count, 1)
                    rl_table = RLTable(
                        table_data,
                        colWidths=[col_width] * col_count,
                    )
                    rl_table.setStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.5, (0, 0, 0)),
                            ("FONTNAME", (0, 0), (-1, -1), cn_font_name),
                            ("FONTSIZE", (0, 0), (-1, -1), 10),
                            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ("TOPPADDING", (0, 0), (-1, -1), 4),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ]
                    )
                    story.append(rl_table)
                story.append(Spacer(1, 8))

            pdf_doc.build(story)

            return {
                "success": True,
                "output_path": pdf_path,
                "method": "reportlab",
            }
        except Exception as e:
            return {"success": False, "error": f"reportlab生成PDF失败: {e}"}

    def _register_chinese_font(self) -> str:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        font_candidates = [
            ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
            ("SimSun", "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            ("SimSun", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
            ("SimHei", r"C:\Windows\Fonts\simhei.ttf"),
            ("SimHei", "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
        ]

        for font_name, font_path in font_candidates:
            if os.path.isfile(font_path):
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    return font_name
                except Exception:
                    continue

        try:
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont

            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            return "STSong-Light"
        except Exception:
            pass

        return "Helvetica"

    @staticmethod
    def _escape_xml(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    def _count_pdf_pages(self, pdf_path: str) -> int:
        try:
            import fitz

            doc = fitz.open(pdf_path)
            count = doc.page_count
            doc.close()
            return count
        except ImportError:
            pass
        except Exception:
            pass

        try:
            with open(pdf_path, "rb") as f:
                content = f.read()
            import re

            matches = re.findall(rb"/Type\s*/Page[^s]", content)
            return max(len(matches), 1)
        except Exception:
            return 0

    @staticmethod
    def _build_bookmark_info(bookmarks: list[dict[str, Any]]) -> list[dict[str, Any]]:
        result = []
        for bm in bookmarks:
            result.append(
                {
                    "title": bm["title"],
                    "level": bm["level"],
                }
            )
        return result
