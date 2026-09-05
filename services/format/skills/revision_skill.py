from __future__ import annotations

import difflib
import logging
import os
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from core.skill_engine.base import Skill, SkillContext, SkillResult

logger = logging.getLogger(__name__)

REVISION_AUTHOR = "TenderPilot"
REVISION_DATE = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")


class RevisionSkill(Skill):
    name = "revision"
    description = "文档版本对比，生成修订追踪标记"
    category = "output"
    version = "1.0.0"
    triggers = ["修订", "版本对比", "修订追踪", "文档对比", "diff"]

    async def execute(self, ctx: SkillContext) -> SkillResult:
        original_path = ctx.parameters.get("original_path", "")
        revised_path = ctx.parameters.get("revised_path", "")
        output_dir = ctx.parameters.get("output_dir", "")
        author = ctx.parameters.get("author", REVISION_AUTHOR)

        if not original_path:
            return SkillResult(success=False, error="未提供原始文件路径(original_path)")
        if not revised_path:
            return SkillResult(success=False, error="未提供修订文件路径(revised_path)")

        if not Path(original_path).exists():
            return SkillResult(success=False, error=f"原始文件不存在: {original_path}")
        if not Path(revised_path).exists():
            return SkillResult(success=False, error=f"修订文件不存在: {revised_path}")

        try:
            from docx import Document
        except ImportError:
            return SkillResult(success=False, error="python-docx库未安装")

        try:
            original_doc = Document(original_path)
            revised_doc = Document(revised_path)
        except Exception as e:
            return SkillResult(success=False, error=f"读取文档失败: {e}")

        original_paras = self._extract_paragraphs(original_doc)
        revised_paras = self._extract_paragraphs(revised_doc)

        diff_ops = self._compute_paragraph_diff(original_paras, revised_paras)

        output_doc = self._build_revision_document(original_doc, revised_doc, diff_ops, author)

        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(
                output_dir,
                Path(original_path).stem + "_revised.docx",
            )
        else:
            output_path = str(Path(original_path).with_suffix("") + "_revised.docx")

        try:
            output_doc.save(output_path)
        except Exception as e:
            return SkillResult(success=False, error=f"保存修订文档失败: {e}")

        added_count = sum(1 for op in diff_ops if op["type"] == "added")
        deleted_count = sum(1 for op in diff_ops if op["type"] == "deleted")
        modified_count = sum(1 for op in diff_ops if op["type"] == "modified")

        return SkillResult(
            success=True,
            data={
                "output_path": output_path,
                "added_count": added_count,
                "deleted_count": deleted_count,
                "modified_count": modified_count,
                "total_changes": added_count + deleted_count + modified_count,
                "original_paragraphs": len(original_paras),
                "revised_paragraphs": len(revised_paras),
            },
        )

    def _extract_paragraphs(self, doc: Any) -> list[dict[str, Any]]:
        paragraphs: list[dict[str, Any]] = []
        for para in doc.paragraphs:
            text = para.text
            style_name = para.style.name if para.style else ""
            paragraphs.append(
                {
                    "text": text,
                    "style": style_name,
                    "element": para._element,
                }
            )
        return paragraphs

    def _compute_paragraph_diff(
        self,
        original_paras: list[dict[str, Any]],
        revised_paras: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        original_texts = [p["text"] for p in original_paras]
        revised_texts = [p["text"] for p in revised_paras]

        matcher = difflib.SequenceMatcher(None, original_texts, revised_texts, autojunk=False)

        diff_ops: list[dict[str, Any]] = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                continue
            elif tag == "replace":
                deleted_items = original_paras[i1:i2]
                added_items = revised_paras[j1:j2]

                if len(deleted_items) == len(added_items):
                    for k in range(len(deleted_items)):
                        old_text = deleted_items[k]["text"]
                        new_text = added_items[k]["text"]
                        if old_text == new_text:
                            continue
                        char_diff = self._compute_char_diff(old_text, new_text)
                        diff_ops.append(
                            {
                                "type": "modified",
                                "original_idx": i1 + k,
                                "revised_idx": j1 + k,
                                "original_text": old_text,
                                "revised_text": new_text,
                                "char_diff": char_diff,
                            }
                        )
                else:
                    for k in range(i1, i2):
                        diff_ops.append(
                            {
                                "type": "deleted",
                                "original_idx": k,
                                "revised_idx": None,
                                "original_text": original_paras[k]["text"],
                                "revised_text": "",
                            }
                        )
                    for k in range(j1, j2):
                        diff_ops.append(
                            {
                                "type": "added",
                                "original_idx": None,
                                "revised_idx": k,
                                "original_text": "",
                                "revised_text": revised_paras[k]["text"],
                            }
                        )
            elif tag == "delete":
                for k in range(i1, i2):
                    diff_ops.append(
                        {
                            "type": "deleted",
                            "original_idx": k,
                            "revised_idx": None,
                            "original_text": original_paras[k]["text"],
                            "revised_text": "",
                        }
                    )
            elif tag == "insert":
                for k in range(j1, j2):
                    diff_ops.append(
                        {
                            "type": "added",
                            "original_idx": None,
                            "revised_idx": k,
                            "original_text": "",
                            "revised_text": revised_paras[k]["text"],
                        }
                    )

        return diff_ops

    @staticmethod
    def _compute_char_diff(old_text: str, new_text: str) -> list[dict[str, Any]]:
        matcher = difflib.SequenceMatcher(None, old_text, new_text, autojunk=False)
        segments: list[dict[str, Any]] = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == "equal":
                segments.append(
                    {
                        "type": "equal",
                        "text": old_text[i1:i2],
                    }
                )
            elif tag == "replace":
                segments.append(
                    {
                        "type": "deleted",
                        "text": old_text[i1:i2],
                    }
                )
                segments.append(
                    {
                        "type": "inserted",
                        "text": new_text[j1:j2],
                    }
                )
            elif tag == "delete":
                segments.append(
                    {
                        "type": "deleted",
                        "text": old_text[i1:i2],
                    }
                )
            elif tag == "insert":
                segments.append(
                    {
                        "type": "inserted",
                        "text": new_text[j1:j2],
                    }
                )

        return segments

    def _build_revision_document(
        self,
        original_doc: Any,
        revised_doc: Any,
        diff_ops: list[dict[str, Any]],
        author: str,
    ) -> Any:
        from docx.oxml.ns import qn

        output_doc = deepcopy(revised_doc)

        self._enable_revision_tracking(output_doc, author)

        body = output_doc.element.body
        existing_paras = list(body.iterchildren(qn("w:p")))

        deleted_ops = [op for op in diff_ops if op["type"] == "deleted"]
        modified_ops = [op for op in diff_ops if op["type"] == "modified"]

        for op in deleted_ops:
            del_para = self._create_deleted_paragraph(op["original_text"], author)
            if existing_paras:
                body.insert(0, del_para)
            else:
                body.append(del_para)

        for op in modified_ops:
            self._apply_modified_paragraph_markup(output_doc, op, author)

        return output_doc

    def _enable_revision_tracking(self, doc: Any, author: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        settings_element = doc.settings.element

        track_changes = settings_element.find(qn("w:trackChanges"))
        if track_changes is None:
            track_changes = OxmlElement("w:trackChanges")
            settings_element.append(track_changes)

    def _create_deleted_paragraph(self, text: str, author: str) -> Any:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p = OxmlElement("w:p")

        p_pr = OxmlElement("w:pPr")
        r_pr_change = OxmlElement("w:rPr")

        del_elem = OxmlElement("w:del")
        del_elem.set(qn("w:id"), str(self._next_id()))
        del_elem.set(qn("w:author"), author)
        del_elem.set(qn("w:date"), REVISION_DATE)

        r_pr_change.append(del_elem)
        p_pr.append(r_pr_change)
        p.append(p_pr)

        if text:
            del_run = OxmlElement("w:del")
            del_run.set(qn("w:id"), str(self._next_id()))
            del_run.set(qn("w:author"), author)
            del_run.set(qn("w:date"), REVISION_DATE)

            r = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            r.append(r_pr)

            t = OxmlElement("w:delText")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)

            del_run.append(r)
            p.append(del_run)

        return p

    def _apply_modified_paragraph_markup(
        self,
        doc: Any,
        op: dict[str, Any],
        author: str,
    ) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        char_diff = op.get("char_diff", [])
        if not char_diff:
            return

        paras = doc.paragraphs
        revised_idx = op.get("revised_idx")
        if revised_idx is None or revised_idx >= len(paras):
            return

        target_para = paras[revised_idx]
        p_element = target_para._element

        for child in list(p_element):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "r":
                p_element.remove(child)

        p_pr = p_element.find(qn("w:pPr"))
        insert_pos = 1 if p_pr is not None else 0

        for segment in char_diff:
            seg_text = segment["text"]
            if not seg_text:
                continue

            if segment["type"] == "equal":
                r = OxmlElement("w:r")
                r_pr = OxmlElement("w:rPr")
                r.append(r_pr)
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = seg_text
                r.append(t)
                p_element.insert(insert_pos, r)
                insert_pos += 1

            elif segment["type"] == "deleted":
                del_elem = OxmlElement("w:del")
                del_elem.set(qn("w:id"), str(self._next_id()))
                del_elem.set(qn("w:author"), author)
                del_elem.set(qn("w:date"), REVISION_DATE)

                r = OxmlElement("w:r")
                r_pr = OxmlElement("w:rPr")
                r.append(r_pr)
                t = OxmlElement("w:delText")
                t.set(qn("xml:space"), "preserve")
                t.text = seg_text
                r.append(t)

                del_elem.append(r)
                p_element.insert(insert_pos, del_elem)
                insert_pos += 1

            elif segment["type"] == "inserted":
                ins_elem = OxmlElement("w:ins")
                ins_elem.set(qn("w:id"), str(self._next_id()))
                ins_elem.set(qn("w:author"), author)
                ins_elem.set(qn("w:date"), REVISION_DATE)

                r = OxmlElement("w:r")
                r_pr = OxmlElement("w:rPr")
                r.append(r_pr)
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = seg_text
                r.append(t)

                ins_elem.append(r)
                p_element.insert(insert_pos, ins_elem)
                insert_pos += 1

    _id_counter = 100

    @classmethod
    def _next_id(cls) -> int:
        cls._id_counter += 1
        return cls._id_counter
