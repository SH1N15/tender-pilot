from __future__ import annotations

import os
import tempfile

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from services.database import get_db
from services.models import (
    Chapter,
    CheckReport,
    CheckType,
    Document,
    Project,
)

router = APIRouter()


class CheckGraphCreate(BaseModel):
    check_ids: list[str] | None = None
    formats: list[str] = ["markdown", "html"]


async def _run_check_graph_compat(project_id: str, check_ids: list[str] | None, db: AsyncSession, timeout: float):
    """旧检查端点兼容垫片：统一走 G-3 CheckRunManager，再映射旧响应。"""
    from services.check.graph_runtime import get_check_run_manager

    project, tender_text, bid_text = await _get_tender_and_bid_text(project_id, db)
    record = await get_check_run_manager().create_run({
        "project_id": project_id,
        "tender_text": tender_text,
        "bid_text": bid_text,
        "check_ids": check_ids,
        "formats": [],
        "project_name": project.name,
    })
    record = await get_check_run_manager().wait_settled(record.run_id, timeout=timeout)
    snap = record.snapshot or {}
    results = snap.get("results")
    if results is None:
        results = snap.get("report", {}).get("results") or {}
    if isinstance(results, list):
        results = {r.get("check_id", str(i)): r for i, r in enumerate(results)}
    if check_ids:
        item = next(iter(results.values()), {})
        return {
            "success": item.get("status") not in ("error", "skipped"),
            "data": item.get("data", {}),
            "error": (item.get("reason") or item.get("error")) if item.get("status") == "error" else None,
            "warnings": item.get("warnings", []),
        }
    return {
        "success": record.status == "completed",
        "data": {
            item_id: {
                "success": item.get("status") not in ("error", "skipped"),
                "data": item.get("data", {}),
                "error": item.get("reason") or item.get("error"),
            }
            for item_id, item in results.items()
        },
        "has_critical": any(item.get("status") in ("fail", "error") for item in results.values()),
    }


async def _persist_graph_report(project_id: str, report: dict) -> dict:
    """图节点使用独立会话落检查报告表，避免持有请求会话。"""
    from services.database import async_session

    async with async_session()() as session:
        row = CheckReport(
            project_id=project_id,
            type="full",
            results=report.get("results") or {},
            risk_level=report.get("risk_level", "low"),
            summary=report.get("summary") or {},
        )
        session.add(row)
        await session.commit()
        return {"report_id": row.id, "type": row.type}


@router.post("/{project_id}/graph")
async def create_check_graph_run(
    project_id: str,
    payload: CheckGraphCreate,
    db: AsyncSession = Depends(get_db),
):
    """单项/全量检查统一图入口；不改变旧直调端点。"""
    from services.check.graph_runtime import get_check_run_manager

    project, tender_text, bid_text = await _get_tender_and_bid_text(project_id, db)
    manager = get_check_run_manager()
    record = await manager.create_run(
        {
            "project_id": project_id,
            "tender_text": tender_text,
            "bid_text": bid_text,
            "check_ids": payload.check_ids,
            "formats": payload.formats,
            "project_name": project.name,
        }
    )
    return {"success": True, "run_id": record.run_id, "status": record.status}


@router.get("/{project_id}/graph/{run_id}")
async def get_check_graph_run(project_id: str, run_id: str):
    """查询图运行快照（报告与 MD/HTML 导出内容均由图节点产出）。"""
    from services.check.graph_runtime import get_check_run_manager

    try:
        record = get_check_run_manager().get(run_id, project_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="检查图 run 不存在")
    return {
        "success": True,
        "run_id": run_id,
        "project_id": record.project_id,
        "status": record.status,
        "error": record.error,
        "snapshot": record.snapshot,
    }


@router.post("/{project_id}/graph/{run_id}/resume")
async def resume_check_graph_run(project_id: str, run_id: str):
    """kill 后从最近 checkpoint 恢复检查图（已完成项不重跑；completed 幂等返回）。"""
    from services.check.graph_runtime import get_check_run_manager

    manager = get_check_run_manager()
    try:
        await manager.resume(run_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="检查图 run 不存在（checkpoint 亦无）")
    try:
        record = manager.get(run_id, project_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="检查图 run 不存在")
    return {
        "success": True,
        "run_id": run_id,
        "project_id": record.project_id,
        "status": record.status,
        "error": record.error,
        "snapshot": record.snapshot,
    }


async def _get_tender_and_bid_text(project_id: str, db: AsyncSession):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    tender_text = ""
    if project.tender_doc_id:
        doc_result = await db.execute(select(Document).where(Document.id == project.tender_doc_id))
        doc = doc_result.scalar_one_or_none()
        if doc and doc.parsed_content:
            tender_text = doc.parsed_content

    bid_docs = await db.execute(
        select(Document).where(
            Document.project_id == project.id,
            Document.type == "bid",
        )
    )
    bid_doc = bid_docs.scalars().first()
    bid_text = bid_doc.parsed_content if bid_doc and bid_doc.parsed_content else ""

    if not bid_text:
        chapter_result = await db.execute(select(Chapter).where(Chapter.project_id == project.id))
        chapters = chapter_result.scalars().all()
        if chapters:
            bid_text = "\n\n".join(f"## {ch.title}\n{ch.content or ''}" for ch in chapters if ch.content)

    return project, tender_text, bid_text


@router.post("/{project_id}/compliance")
async def check_compliance(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["compliance_check"], db, 90)


@router.post("/{project_id}/disqualification")
async def check_disqualification(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["disqualification_check"], db, 90)


@router.post("/{project_id}/qualification")
async def check_qualification(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["qualification_check"], db, 90)


@router.post("/{project_id}/pricing")
async def check_pricing(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["pricing_check"], db, 90)


@router.post("/{project_id}/fit-score")
async def check_fit_score(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["fit_score"], db, 90)


@router.post("/{project_id}/selfcheck")
async def run_selfcheck(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["selfcheck_list"], db, 90)


@router.post("/{project_id}/full-check")
async def full_check(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, None, db, 420)


@router.get("/{project_id}/reports")
async def list_check_reports(project_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    reports_result = await db.execute(select(CheckReport).where(CheckReport.project_id == project.id))
    reports = reports_result.scalars().all()

    return {
        "reports": [
            {
                "id": str(r.id),
                "type": r.type.value if isinstance(r.type, CheckType) else r.type,
                "risk_level": r.risk_level,
                "summary": r.summary,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in reports
        ]
    }


@router.post("/{project_id}/deposit")
async def check_deposit(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["deposit_check"], db, 90)


async def _parse_uploaded_file(file: UploadFile) -> str:
    suffix = os.path.splitext(file.filename or "")[1].lower()
    content_bytes = await file.read()

    if suffix == ".txt" or suffix == ".md":
        return content_bytes.decode("utf-8", errors="replace")

    if suffix == ".docx":
        tmp_path = ""
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name

            from docx import Document as DocxDocument

            doc = DocxDocument(tmp_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style_name = (para.style.name or "").lower() if para.style else ""
                    if "heading" in style_name or "标题" in style_name:
                        level = "1"
                        for ch in style_name:
                            if ch.isdigit():
                                level = ch
                                break
                        paragraphs.append(f"{'#' * int(level)} {text}")
                    else:
                        paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        paragraphs.append(" | ".join(cells))

            return "\n\n".join(paragraphs)
        except Exception:
            return content_bytes.decode("utf-8", errors="replace")
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    if suffix == ".pdf":
        tmp_path = ""
        try:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name

            try:
                import pdfplumber

                texts = []
                with pdfplumber.open(tmp_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            texts.append(page_text)
                return "\n\n".join(texts)
            except ImportError:
                pass

            try:
                import fitz

                doc = fitz.open(tmp_path)
                texts = []
                for page in doc:
                    texts.append(page.get_text())
                doc.close()
                return "\n\n".join(texts)
            except ImportError:
                pass

            return content_bytes.decode("utf-8", errors="replace")
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    return content_bytes.decode("utf-8", errors="replace")


# G-5：upload-check 图化后的映射表（旧 _CHECK_SKILL_MAP 直调表已删）
# camelCase（前端 check_type）→ 图 check_id（snake_case，同 services/check/graph_adapter.py CHECK_REGISTRY）
_CHECK_TYPE_TO_GRAPH_ID = {
    "compliance": "compliance_check",
    "disqualification": "disqualification_check",
    "qualification": "qualification_check",
    "pricing": "pricing_check",
    "fitScore": "fit_score",
    "deposit": "deposit_check",
    "signature": "signature_check",
    "validity": "validity_check",
    "consistency": "consistency_check",
    "duplicate": "duplicate_check",
    "mandatoryReq": "mandatory_req_check",
    "docIntegrity": "doc_integrity_check",
    "aiTextCheck": "ai_text_check",
    "riskScore": "risk_score",
    "crossCheck": "cross_check",
    "sampleReport": "sample_report_check",
    "jointBid": "joint_bid_check",
    "ebidSubmit": "ebid_submit_check",
    "pricingLogic": "pricing_logic_check",
    "selfcheck": "selfcheck_list",
}
# check_id → camelCase（响应键名回映射；whitelist_filter 等非前端项自动出局）
_GRAPH_CHECK_IDS = {v: k for k, v in _CHECK_TYPE_TO_GRAPH_ID.items()}
# 旧 fullCheck 上传检查的 15 项集合（与旧 check_types 列表一致）
_UPLOAD_FULL_CHECK_IDS = [
    _CHECK_TYPE_TO_GRAPH_ID[c]
    for c in (
        "compliance",
        "disqualification",
        "qualification",
        "pricing",
        "fitScore",
        "deposit",
        "signature",
        "validity",
        "consistency",
        "duplicate",
        "mandatoryReq",
        "docIntegrity",
        "aiTextCheck",
        "crossCheck",
        "pricingLogic",
    )
]


@router.post("/upload-check")
async def upload_and_check(
    bid_file: UploadFile = File(..., description="投标文件(.docx/.pdf/.txt)"),
    tender_file: UploadFile | None = File(None, description="招标文件(可选，.docx/.pdf/.txt)"),
    check_type: str = Form("fullCheck", description="检查类型: fullCheck/compliance/disqualification/..."),
    db: AsyncSession = Depends(get_db),
):
    """上传检查：文本解析后统一走 G-3 检查图执行（G-5 图化，响应形状与旧直调一致）。

    - project_id 用伪 ID（upload_ 前缀），persist_report=False 不落检查报告表；
    - 单项/全量的响应映射与 _run_check_graph_compat 同口径。
    """
    import uuid as _uuid

    from services.check.graph_runtime import get_check_run_manager

    bid_text = await _parse_uploaded_file(bid_file)
    if not bid_text.strip():
        raise HTTPException(status_code=400, detail="投标文件内容为空或无法解析")

    tender_text = ""
    if tender_file:
        tender_text = await _parse_uploaded_file(tender_file)

    if check_type == "fullCheck":
        graph_check_ids = _UPLOAD_FULL_CHECK_IDS
    else:
        graph_check_ids = [_CHECK_TYPE_TO_GRAPH_ID.get(check_type, check_type)]
        if graph_check_ids[0] not in _GRAPH_CHECK_IDS:
            raise HTTPException(status_code=400, detail=f"不支持的检查类型: {check_type}")

    manager = get_check_run_manager()
    record = await manager.create_run(
        {
            "project_id": f"upload_{_uuid.uuid4().hex[:8]}",
            "tender_text": tender_text,
            "bid_text": bid_text,
            "check_ids": graph_check_ids,
            "formats": [],
            "project_name": "上传检查",
            "persist_report": False,
        }
    )
    record = await manager.wait_settled(record.run_id, timeout=900)
    if record.status == "failed":
        raise HTTPException(status_code=500, detail=record.error or "检查图执行失败")
    snap = record.snapshot or {}
    results = snap.get("results")
    if results is None:
        results = snap.get("report", {}).get("results") or {}
    if isinstance(results, list):
        results = {r.get("check_id", str(i)): r for i, r in enumerate(results)}
    # 图 check_id(snake) → 前端 camelCase 键（保留旧响应键名）
    by_camel = {_GRAPH_CHECK_IDS.get(k, k): v for k, v in results.items()}

    def _item_shape(item: dict) -> dict:
        return {
            "success": item.get("status") not in ("error", "skipped"),
            "data": item.get("data", {}),
            "error": (item.get("reason") or item.get("error")) if item.get("status") == "error" else None,
            "warnings": item.get("warnings", []),
        }

    if check_type == "fullCheck":
        all_results = {camel: _item_shape(item) for camel, item in by_camel.items()}
        has_critical = any(
            r.get("data", {}).get("has_critical_issues") or r.get("data", {}).get("risk_level") == "high"
            for r in all_results.values()
            if r.get("success") and isinstance(r.get("data"), dict)
        )
        return {
            "success": True,
            "data": all_results,
            "has_critical": has_critical,
            "source": "upload",
            "bid_filename": bid_file.filename,
            "tender_filename": tender_file.filename if tender_file else None,
        }

    item = next(iter(by_camel.values()), {})
    shaped = _item_shape(item)
    return {
        "success": shaped["success"],
        "data": shaped["data"],
        "error": shaped["error"],
        "warnings": shaped["warnings"],
        "source": "upload",
        "bid_filename": bid_file.filename,
        "tender_filename": tender_file.filename if tender_file else None,
    }


@router.post("/{project_id}/signature")
async def check_signature(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["signature_check"], db, 90)


@router.post("/{project_id}/validity")
async def check_validity(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["validity_check"], db, 90)


@router.post("/{project_id}/consistency")
async def check_consistency(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["consistency_check"], db, 90)


@router.post("/{project_id}/duplicate")
async def check_duplicate(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["duplicate_check"], db, 90)


@router.post("/{project_id}/mandatory-req")
async def check_mandatory_req(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["mandatory_req_check"], db, 90)


@router.get("/{project_id}/reports/{report_id}/export")
async def export_check_report(
    project_id: str,
    report_id: str,
    format: str = "markdown",
    db: AsyncSession = Depends(get_db),
):
    from fastapi.responses import PlainTextResponse

    from services.check.report_export import render_check_report_export

    result = await db.execute(select(CheckReport).where(CheckReport.id == report_id))
    report = result.scalar_one_or_none()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    project_result = await db.execute(select(Project).where(Project.id == project_id))
    project = project_result.scalar_one_or_none()
    project_name = project.name if project else "未命名项目"

    # G-5：导出为确定性格式化（无 LLM），执行体在 services/check/report_export.py
    try:
        content = await render_check_report_export(report.results or {}, format, project_name)
    except ValueError as e:
        raise HTTPException(status_code=500, detail=str(e))
    content_type = "text/markdown" if format == "markdown" else "text/html" if format == "html" else "application/json"
    return PlainTextResponse(content=content, media_type=content_type)


@router.post("/{project_id}/doc-integrity")
async def check_doc_integrity(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["doc_integrity_check"], db, 90)


@router.post("/{project_id}/ai-text-check")
async def check_ai_text(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["ai_text_check"], db, 90)


@router.post("/{project_id}/risk-score")
async def check_risk_score(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["risk_score"], db, 90)


@router.post("/{project_id}/cross-check")
async def check_cross(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["cross_check"], db, 90)


@router.post("/{project_id}/sample-report")
async def check_sample_report(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["sample_report_check"], db, 90)


@router.post("/{project_id}/joint-bid")
async def check_joint_bid(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["joint_bid_check"], db, 90)


@router.post("/{project_id}/ebid-submit")
async def check_ebid_submit(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["ebid_submit_check"], db, 90)


@router.post("/{project_id}/pricing-logic")
async def check_pricing_logic(project_id: str, db: AsyncSession = Depends(get_db)):
    # G-5 T3：return 之后的不可达旧直调代码已删除
    return await _run_check_graph_compat(project_id, ["pricing_logic_check"], db, 90)


@router.get("/{project_id}/missing-materials-docx")
async def export_missing_materials_docx(
    project_id: str,
    format: str = "docx",
    db: AsyncSession = Depends(get_db),
):
    """Worker I 任务2：《需补充材料清单》导出（docx 表格 / markdown，确定性提取无 LLM）。

    数据两路合并：最新全量检查报告 results 的 fact 型缺失 finding + 章节正文
    【待补充】/（知识库无据）占位标注；按 (章节, 指纹) 去重、优先级排序。
    """
    from urllib.parse import quote

    from fastapi import Response
    from fastapi.responses import PlainTextResponse

    from services.check.missing_materials import (
        build_missing_materials_docx,
        build_missing_materials_items,
        render_missing_materials_markdown,
    )

    fmt = (format or "docx").lower()
    if fmt not in ("docx", "md", "markdown"):
        raise HTTPException(status_code=400, detail="format 仅支持 docx / markdown")

    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    report_result = await db.execute(
        select(CheckReport)
        .where(CheckReport.project_id == project_id)
        .order_by(CheckReport.created_at.desc())
        .limit(1)
    )
    latest_report = report_result.scalars().first()

    chapter_result = await db.execute(
        select(Chapter.id, Chapter.title, Chapter.content).where(Chapter.project_id == project_id)
    )
    chapters = [(str(r.id), str(r.title or ""), str(r.content or "")) for r in chapter_result.all()]

    items = build_missing_materials_items(latest_report.results if latest_report else None, chapters)
    filename = quote(f"{project.name}-需补充材料清单")

    if fmt in ("md", "markdown"):
        return PlainTextResponse(
            content=render_missing_materials_markdown(project.name, items),
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}.md"},
        )
    docx_bytes = build_missing_materials_docx(project.name, items)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}.docx"},
    )
