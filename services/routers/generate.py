from __future__ import annotations

import asyncio
import io
import json
import logging

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from core.task_manager import TaskManager
from services.database import get_db
from services.llm_factory import get_llm_gateway
from services.models import Chapter, Document, Outline, Project

logger = logging.getLogger(__name__)
router = APIRouter()


class IllustrationGenerateRequest(BaseModel):
    provider: str | None = None
    image_size: str = "landscape_16_9"
    regenerate: bool = False


# ─────────────────────────────────────────────
# 异步任务端点
# ─────────────────────────────────────────────


# ─────────────────────────────────────────────
# P-G G-2：章节生成图模式（触发器/查询器；直调端点原样保留，G-4 收口）
# ─────────────────────────────────────────────


class GraphRunCreate(BaseModel):
    outline_mode: str = "aligned"
    run_outline: bool = True
    outline_only: bool = False  # True=只跑大纲（前端大纲入口用，防大项目全量正文 runaway）
    chapter_modes: dict[str, str] | None = None  # {chapter_id: "A"|"B"}，缺省全 B
    chapter_ids: list[str] | None = None  # 限定章节，缺省全部


class OutlineUpdate(BaseModel):
    tree: dict | list


@router.post("/{project_id}/graph/run")
async def create_graph_run(project_id: str, payload: GraphRunCreate, db: AsyncSession = Depends(get_db)):
    """创建章节生成图运行（大纲→逐章正文→Grounding 硬门→落库，逐章 checkpoint）。"""
    from services.generate.graph_runtime import get_generation_run_manager

    result = await db.execute(select(Project).where(Project.id == project_id))
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="项目不存在")
    manager = get_generation_run_manager()
    record = await manager.create_run(payload.model_dump() | {"project_id": project_id})
    return {"success": True, "run_id": record.run_id, "status": record.status}


@router.get("/{project_id}/graph/runs")
async def list_graph_runs(project_id: str):
    from services.generate.graph_runtime import get_generation_run_manager

    return {
        "success": True,
        "runs": [
            {
                "run_id": r.run_id,
                "project_id": r.project_id,
                "status": r.status,
                "created_at": r.created_at,
                "error": r.error,
            }
            for r in get_generation_run_manager().list_runs(project_id)
        ],
    }


@router.get("/{project_id}/graph/runs/{run_id}")
async def get_graph_run(project_id: str, run_id: str):
    from services.generate.graph_runtime import get_generation_run_manager

    try:
        record = get_generation_run_manager().get(run_id, project_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="run 不存在")
    manager = get_generation_run_manager()
    # 运行中：读 checkpointer 实时快照（逐章进度轮询）
    snap = record.snapshot if record.snapshot else await manager.live_snapshot(run_id, project_id)
    return {"success": True, **manager.snapshot_payload(record, snap)}


@router.post("/{project_id}/graph/runs/{run_id}/resume")
async def resume_graph_run(project_id: str, run_id: str):
    """kill 后从最近 checkpoint 续写（已完成章不重写；completed run 幂等返回）。"""
    from services.generate.graph_runtime import get_generation_run_manager

    manager = get_generation_run_manager()
    try:
        await manager.resume(run_id, project_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="run 不存在")
    # 等 resume 后台任务跑到终态（供脚本轮询之外的一次性同步观察；超时不阻塞）
    try:
        await manager.wait_settled(run_id, timeout=1.0)
    except (KeyError, asyncio.TimeoutError):
        pass
    except Exception:  # noqa: BLE001
        pass
    try:
        record = manager.get(run_id, project_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="run 不存在")
    return {"success": True, **manager.snapshot_payload(record)}


@router.get("/task/{task_id}")
async def get_task_status(task_id: str):
    """查询异步任务状态"""
    tm = TaskManager.instance()
    task = tm.get_task(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return task.to_dict()


# ─────────────────────────────────────────────
# 大纲生成（异步任务模式）
# ─────────────────────────────────────────────


def _iter_outline_chapters(nodes):
    """深度优先遍历大纲树的章节节点，产出 (node_id, title) 元组。"""
    for node in nodes:
        if not isinstance(node, dict):
            continue
        node_id = node.get("id")
        if node_id is None or not node.get("title"):
            continue
        yield str(node_id), str(node["title"])
        children = node.get("children", [])
        if isinstance(children, list):
            yield from _iter_outline_chapters(children)


async def _do_generate_outline(project_id: str, mode: str):
    """大纲生成的实际执行逻辑（在后台任务中运行）。

    G-5：执行体迁至 services/generate/direct_exec.py（路由文件直调 Skill 归零）；
    Skill 类与 gateway 仍按原注入方式（tests monkeypatch Skill 类方法不受影响）。
    """
    from services.generate.direct_exec import do_generate_outline_direct

    return await do_generate_outline_direct(project_id, mode, get_llm_gateway())


async def _do_generate_outline_graph(project_id: str, mode: str):
    """G-4R HTTP 垫片使用的图执行函数；保留旧 helper 供内部回归测试。"""
    from services.generate.graph_runtime import get_generation_run_manager

    record = await get_generation_run_manager().create_run({
        "project_id": project_id, "outline_mode": mode, "run_outline": True,
        "chapter_ids": [], "chapter_modes": {},
    })
    record = await get_generation_run_manager().wait_settled(record.run_id, timeout=420)
    snap = record.snapshot or {}
    return {
        "success": record.status == "completed",
        "data": {
            "outline": snap.get("outline") or snap.get("outline_tree") or {},
            "chapters": snap.get("chapters", []),
        },
        "error": record.error or (snap.get("errors") or [None])[0],
        "warnings": snap.get("warnings", []),
    }


@router.post("/{project_id}/outline")
async def generate_outline(
    project_id: str,
    mode: str = "aligned",
    db: AsyncSession = Depends(get_db),
):
    """大纲生成（异步任务模式）。

    立即返回 task_id，前端通过 GET /generate/task/{task_id} 轮询结果。
    兼容模式：如果请求带 ?sync=1 则走同步模式（调试用）。
    """
    # 快速校验
    result = await db.execute(select(Project).where(Project.id == project_id))
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="项目不存在")

    # 提交异步任务
    tm = TaskManager.instance()
    task = await tm.submit("outline_gen", _do_generate_outline_graph, project_id, mode)

    return {
        "task_id": task.task_id,
        "status": "pending",
        "message": "大纲生成任务已提交，请通过 GET /generate/task/{task_id} 查询进度",
    }


def _flatten_outline_nodes(nodes: list[dict], depth: int = 1):
    """Flatten the editable outline while retaining DFS order for chapter materialization."""
    rows: list[tuple[str, str, int, int]] = []
    for index, raw in enumerate(nodes):
        if not isinstance(raw, dict):
            continue
        node_id = str(raw.get("id") or "").strip()
        title = str(raw.get("title") or "").strip()
        if not node_id or not title:
            continue
        try:
            level = int(raw.get("level") or depth)
        except (TypeError, ValueError):
            level = depth
        rows.append((node_id, title, max(level, 1), len(rows)))
        children = raw.get("children")
        if isinstance(children, list):
            child_rows = _flatten_outline_nodes(children, depth + 1)
            rows.extend((cid, ctitle, clevel, len(rows)) for cid, ctitle, clevel, _ in child_rows)
    return rows


@router.put("/{project_id}/outline")
async def update_outline(project_id: str, payload: OutlineUpdate, db: AsyncSession = Depends(get_db)):
    """Persist edits made in the advanced outline tool and reconcile materialized chapters."""
    project_result = await db.execute(select(Project).where(Project.id == project_id))
    project = project_result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    tree = payload.tree if isinstance(payload.tree, dict) else {"chapters": payload.tree}
    raw_nodes = tree.get("chapters") if isinstance(tree, dict) else []
    if not isinstance(raw_nodes, list):
        raise HTTPException(status_code=422, detail="大纲章节格式无效")
    rows = _flatten_outline_nodes(raw_nodes)
    if not rows:
        raise HTTPException(status_code=422, detail="大纲至少需要一个有效章节")

    outline_result = await db.execute(select(Outline).where(Outline.project_id == project_id))
    outline = outline_result.scalar_one_or_none()
    if outline:
        outline.tree = tree
    else:
        outline = Outline(project_id=project_id, mode="aligned", tree=tree, score_mapping={})
        db.add(outline)
        await db.flush()

    chapter_result = await db.execute(select(Chapter).where(Chapter.project_id == project_id))
    existing = {chapter.id: chapter for chapter in chapter_result.scalars().all()}
    ids = {node_id for node_id, _, _, _ in rows}
    for node_id, title, level, sort_order in rows:
        chapter = existing.get(node_id)
        if chapter:
            chapter.title = title
            chapter.sort_order = sort_order
        else:
            db.add(Chapter(
                id=node_id,
                project_id=project_id,
                outline_id=outline.id,
                title=title,
                status="pending",
                content="",
                sort_order=sort_order,
            ))
    removed_ids = set(existing) - ids
    if removed_ids:
        await db.execute(delete(Chapter).where(Chapter.project_id == project_id, Chapter.id.in_(removed_ids)))
    await db.commit()
    return {"success": True, "chapters": len(rows), "removed": len(removed_ids)}


@router.post("/{project_id}/structure/{structure_type}")
async def generate_structure(
    project_id: str,
    structure_type: str,
    db: AsyncSession = Depends(get_db),
):
    # G-5：执行体迁至 services/generate/direct_exec.py（路由文件直调 Skill 归零）
    from services.generate.direct_exec import run_structure_template

    return await run_structure_template(project_id, structure_type, db, get_llm_gateway())


@router.get("/{project_id}/score-coverage")
async def get_score_coverage(
    project_id: str,
    db: AsyncSession = Depends(get_db),
):
    # G-5：执行体迁至 services/generate/direct_exec.py（路由文件直调 Skill 归零）
    from services.generate.direct_exec import run_score_coverage

    return await run_score_coverage(project_id, db, get_llm_gateway())


def _chapter_natural_key(chapter_id: str) -> tuple[int, ...]:
    """章节节点号自然排序："1" -> (1,)，"1.1" -> (1, 1)，"2.10" -> (2, 10)。

    非数字段（如 "new_1730"）退化为 0，保证排序稳定不抛错。
    """
    parts: list[int] = []
    for seg in str(chapter_id).split("."):
        try:
            parts.append(int(seg))
        except ValueError:
            parts.append(0)
    return tuple(parts)


@router.get("/{project_id}/chapters")
async def list_chapters(
    project_id: str,
    db: AsyncSession = Depends(get_db),
):
    """按大纲树顺序返回项目章节（BUG-16：正文生成章节下拉数据源）。

    返回 [{id, title, level, status, word_count, has_content}]，
    has_content=True 表示该章节已生成过正文。
    """
    result = await db.execute(select(Project).where(Project.id == project_id))
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="项目不存在")

    chapter_result = await db.execute(select(Chapter).where(Chapter.project_id == project_id))
    chapters = chapter_result.scalars().all()
    # 按 sort_order 优先、节点号自然排序兜底（物化顺序即大纲 DFS 顺序）
    chapters.sort(key=lambda c: (c.sort_order or 0, _chapter_natural_key(c.id)))

    return {
        "chapters": [
            {
                "id": c.id,
                "title": c.title,
                "level": str(c.id).count(".") + 1,
                "status": c.status,
                "word_count": c.word_count or 0,
                "has_content": bool(c.content),
                # G-2：引用对照表随行返回（正文页【n】点查来源联动）
                "citation_ledger": c.citation_ledger or None,
            }
            for c in chapters
        ]
    }


@router.get("/{project_id}/chapters/{chapter_id}")
async def get_chapter_content(
    project_id: str,
    chapter_id: str,
    db: AsyncSession = Depends(get_db),
):
    """章节正文 + 引用对照表（G-2：正文页【n】点查来源的数据源）。"""
    result = await db.execute(select(Project).where(Project.id == project_id))
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="项目不存在")
    chapter_result = await db.execute(select(Chapter).where(Chapter.id == chapter_id, Chapter.project_id == project_id))
    chapter = chapter_result.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="章节不存在")
    # Older generated chapters stored citation ledgers as a list. Normalize
    # that legacy shape at the API boundary so the current frontend can open
    # any chapter without a 500, while preserving the underlying data.
    raw_ledger = chapter.citation_ledger or {}
    if isinstance(raw_ledger, dict):
        ledger = raw_ledger
    elif isinstance(raw_ledger, list):
        ledger = {
            str(index + 1): (entry if isinstance(entry, dict) else {"source": str(entry)})
            for index, entry in enumerate(raw_ledger)
        }
    else:
        ledger = {}
    return {
        "chapter_id": chapter.id,
        "title": chapter.title,
        "mode": chapter.mode,
        "status": chapter.status,
        "word_count": chapter.word_count or 0,
        "content": chapter.content or "",
        "citation_ledger": ledger,
        "sources": [
            {
                "n": int(n) if str(n).isdigit() else n,
                "chunk_id": (e.get("chunk_id", "") if isinstance(e, dict) else ""),
                "source": (e.get("source", "") if isinstance(e, dict) else str(e)),
            }
            for n, e in sorted(ledger.items(), key=lambda kv: str(kv[0]))
        ],
    }


@router.post("/{project_id}/chapters/{chapter_id}/illustrations")
async def generate_chapter_illustrations(
    project_id: str,
    chapter_id: str,
    req: IllustrationGenerateRequest | None = None,
    db: AsyncSession = Depends(get_db),
):
    """Generate and persist one chapter's illustrations without exporting a document.

    The generated Markdown image references are written to the chapter content and
    mirrored in ``citation_ledger._illustrations_generated`` so a later whole-document
    export only embeds saved images and does not call the image service again.
    """
    result = await db.execute(select(Project).where(Project.id == project_id))
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="项目不存在")
    chapter_result = await db.execute(
        select(Chapter).where(Chapter.id == chapter_id, Chapter.project_id == project_id)
    )
    chapter = chapter_result.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="章节不存在")
    if not (chapter.content or "").strip():
        raise HTTPException(status_code=400, detail="请先生成本章正文")

    req = req or IllustrationGenerateRequest()
    ledger = dict(chapter.citation_ledger or {})
    existing = ledger.get("_illustrations_generated") or []
    if existing and not req.regenerate:
        from core.agent_engine.illustration_assembly import attach_stored_illustrations

        attached = attach_stored_illustrations(chapter.content or "", existing)
        if attached != (chapter.content or ""):
            chapter.content = attached
            chapter.word_count = len(attached)
            chapter.citation_ledger = ledger
            await db.commit()
        return {"success": True, "chapter_id": chapter.id, "generated": len(existing), "reused": True}

    from core.agent_engine.illustration_assembly import assemble_illustrations, get_illustration_params

    params = get_illustration_params(force_enabled=True, provider_override=req.provider)
    if params is None:
        raise HTTPException(status_code=400, detail="图片服务未配置")
    suggestions = ledger.get("_illustrations") or []
    new_text, images, status = await assemble_illustrations(
        project_id,
        chapter.content or "",
        list(suggestions),
        params,
        image_size=req.image_size,
    )
    if status != "assembled" or not images:
        raise HTTPException(status_code=502, detail="本章未生成可用配图")

    ledger["_illustrations_generated"] = images
    chapter.content = new_text
    chapter.citation_ledger = ledger
    await db.commit()
    return {
        "success": True,
        "chapter_id": chapter.id,
        "generated": len(images),
        "reused": False,
        "images": images,
    }


@router.post("/{project_id}/content/{chapter_id}")
async def generate_chapter(
    project_id: str,
    chapter_id: str,
    mode: str = "A",
    db: AsyncSession = Depends(get_db),
):
    # G-5：执行体迁至 services/generate/direct_exec.py（路由文件直调 Skill 归零）
    from services.generate.direct_exec import run_content_chapter

    return await run_content_chapter(project_id, chapter_id, mode, db, get_llm_gateway())


@router.post("/{project_id}/content/stream/{chapter_id}")
async def stream_generate_chapter(
    project_id: str,
    chapter_id: str,
    mode: str = "A",
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    chapter_result = await db.execute(select(Chapter).where(Chapter.id == chapter_id, Chapter.project_id == project.id))
    chapter = chapter_result.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="章节不存在")

    doc_result = await db.execute(select(Document).where(Document.id == project.tender_doc_id))
    doc = doc_result.scalar_one_or_none()
    tender_context = doc.parsed_content[:4000] if doc and doc.parsed_content else ""

    outline_result = await db.execute(select(Outline).where(Outline.project_id == project.id))
    outline = outline_result.scalar_one_or_none()
    chapter_outline = json.dumps(outline.tree, ensure_ascii=False) if outline and outline.tree else ""

    gateway = get_llm_gateway()

    # P3：B 模式接入引用管线——检索知识库并构建编号引用对照表（与非流式 content_gen_skill B 模式同源）
    ledger: dict = {}
    if mode.upper() == "B":
        try:
            from core.rag_engine.kb_adapter import build_default_knowledge_base

            kb = await build_default_knowledge_base()
            docs = await kb.retrieve(query=f"{chapter.title} {chapter_outline}", top_k=5) if kb else []
            if docs:
                from core.agent_engine.evidence_gate import build_ledger

                ledger = build_ledger(docs)
        except Exception as e:  # noqa: BLE001
            logger.warning("B 模式流式引用管线检索失败，回退 A 模式提示: %s", e)
            ledger = {}

    if ledger:
        materials = "\n\n".join(
            f"【引用{entry['n']}】(chunk_id={entry['chunk_id']}, source={entry['source']}):\n{entry['text'][:800]}"
            for _, entry in sorted(ledger.items())
        )
        messages = [
            {
                "role": "system",
                "content": (
                    f'你是标书撰写专家。基于提供的编号参考材料撰写"{chapter.title}"章节。\n'
                    "要求：\n"
                    "1. 整合参考材料，改写为适合本项目的表述，不得直接复制；\n"
                    "2. 引用锚点规则：凡内容来自某条参考材料的，在该句末尾标注其编号标记（如【1】、【2】）；\n"
                    "   编号只能取参考材料列表中存在的编号，不得编造；\n"
                    "3. 硬事实纪律：金额/日期时限/资质编号/技术参数值等硬事实必须来自参考材料且有对应引用标记；\n"
                    '   库内无据的内容明确写"（知识库无据，待补充）"，绝不编造数值；\n'
                    "4. 内容必须针对本项目，不得使用通用模板套话。直接输出正文文本，不要 JSON。"
                ),
            },
            {
                "role": "user",
                "content": f"章节大纲：{chapter_outline}\n\n编号参考材料：\n{materials[:8000]}",
            },
        ]
    else:
        messages = [
            {
                "role": "system",
                "content": f'你是标书撰写专家。请撰写"{chapter.title}"章节。内容必须针对本项目，不得使用通用模板套话。',
            },
            {
                "role": "user",
                "content": f"招标要求上下文：\n{tender_context[:3000]}",
            },
        ]

    async def event_generator():
        collected_content = []
        try:
            async for chunk in gateway.stream_chat(messages=messages, temperature=0.4 if ledger else 0.5):
                collected_content.append(chunk)
                yield f"data: {json.dumps({'content': chunk}, ensure_ascii=False)}\n\n"

            full_content = "".join(collected_content)
            chapter.content = full_content
            chapter.mode = mode
            chapter.status = "generated"
            chapter.word_count = len(full_content)
            await db.commit()

            final_event: dict = {"done": True, "word_count": len(full_content)}
            if ledger:
                # P3：流式路径同样产出引用账本/grounding 统计（末尾事件携带，最小实现）
                from core.agent_engine.evidence_gate import (
                    ground_hard_facts,
                    ledger_for_output,
                    ledger_texts,
                    make_ledger_anchor_func,
                )

                grounding = ground_hard_facts(full_content, ledger)
                final_event["citation_ledger"] = ledger_for_output(ledger)
                final_event["grounding"] = grounding["stats"]
                final_event["sources"] = [
                    {"n": entry["n"], "chunk_id": entry["chunk_id"], "source": entry["source"]}
                    for _, entry in sorted(ledger.items())
                ]
                try:
                    from eval.metrics.citation import citation_valid_rate

                    final_event["citation_rate"] = citation_valid_rate(
                        grounding["text"],
                        ledger_texts(ledger),
                        anchor_func=make_ledger_anchor_func(ledger),
                    )
                except Exception as e:  # noqa: BLE001
                    logger.warning("流式引用有效率计算失败: %s", e)
            yield f"data: {json.dumps(final_event, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@router.post("/{project_id}/mandatory-extract")
async def extract_mandatory_requirements(
    project_id: str,
    db: AsyncSession = Depends(get_db),
):
    # G-5：执行体迁至 services/generate/direct_exec.py（路由文件直调 Skill 归零）
    from services.generate.direct_exec import run_mandatory_extract

    return await run_mandatory_extract(project_id, db, get_llm_gateway())


# ─────────────────────────────────────────────
# Worker G：章节正文 → DOCX 直连导出（不经排版管线）
# ─────────────────────────────────────────────


def _build_docx_from_chapters(
    project_name: str,
    chapters: list[Chapter],
    contents: dict[str, str] | None = None,
) -> tuple[bytes, list[tuple[str, str, str]]]:
    """按大纲层级顺序把已生成章节正文组装为 DOCX 字节流。

    层级映射：1 / 1.1 / 1.1.1 → Heading 1/2/3（更深层级顺延，上限 Heading 6）。
    contents：Worker I 任务3——{chapter_id: 装配后正文}（配图引用以 Markdown 段落进入
    正文，随后由 embed_images_into_docx 嵌入为图片）；缺省用章节原始正文。
    返回 (docx_bytes, skipped)；skipped = [(chapter_id, title, status)]（空章节/未生成）。
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(project_name, level=0)
    skipped: list[tuple[str, str, str]] = []
    for c in chapters:
        raw = (contents or {}).get(str(c.id)) if contents else None
        content = (raw if raw is not None else (c.content or "")).strip()
        if not content:
            skipped.append((str(c.id), c.title, c.status or ""))
            continue
        depth = min(str(c.id).count(".") + 1, 6)
        doc.add_heading(f"{c.id} {c.title}", level=depth)
        for para_text in content.split("\n"):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
    if skipped:
        doc.add_heading("未生成章节清单", level=1)
        doc.add_paragraph(f"共 {len(skipped)} 章未导出正文：")
        for cid, title, status in skipped:
            doc.add_paragraph(f"{cid} {title}（status={status}）", style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), skipped


async def _assemble_export_illustrations(
    project_id: str,
    chapters: list[Chapter],
    force_enabled: bool = False,
    provider_override: str | None = None,
    image_size: str = "landscape_16_9",
    chapter_id: str | None = None,
) -> dict[str, str]:
    """Legacy export compatibility: assemble image suggestions for old API callers.

    The desktop workflow does not use this path. It generates and persists images
    per chapter first, then exports with ``stored_illustrations_only=true``.
    """
    from core.agent_engine.illustration_assembly import assemble_illustrations, get_illustration_params

    params = get_illustration_params(force_enabled=force_enabled, provider_override=provider_override)
    if params is None:
        return {}
    contents: dict[str, str] = {}
    for c in chapters:
        if chapter_id and str(c.id) != str(chapter_id):
            continue
        content = (c.content or "").strip()
        if not content:
            continue
        ledger = c.citation_ledger if isinstance(c.citation_ledger, dict) else {}
        suggestions = ledger.get("_illustrations") or []
        if not suggestions and "[插图位置" not in content:
            continue
        try:
            new_text, images, status = await assemble_illustrations(
                project_id, content, list(suggestions), params, image_size=image_size
            )
        except Exception as e:  # noqa: BLE001
            logger.warning("直连导出配图装配失败（章节 %s，降级无图）: %s", c.id, e)
            continue
        if status == "assembled" and images:
            contents[str(c.id)] = new_text
    return contents


async def _export_fill_map(db: AsyncSession, project_id: str) -> dict[str, str]:
    """Worker J：从解读结果构建导出占位符回填表（[采购人名称]→买方单位名）。"""
    from core.agent_engine.export_sanitizer import build_export_fill_map
    from services.models import Analysis

    try:
        result = await db.execute(select(Analysis).where(Analysis.project_id == project_id))
        analysis = result.scalar_one_or_none()
    except Exception as e:  # noqa: BLE001
        logger.warning("导出净化回填表构建失败（降级不回填）: %s", e)
        return {}
    return build_export_fill_map(analysis.dimensions if analysis else None)


@router.get("/{project_id}/export-docx")
async def export_docx_direct(
    project_id: str,
    illustrations: bool = False,
    force_illustrations: bool = False,
    illustration_provider: str | None = None,
    illustration_size: str = "landscape_16_9",
    illustration_chapter_id: str | None = None,
    stored_illustrations_only: bool = True,
    db: AsyncSession = Depends(get_db),
):
    """章节正文直连导出 DOCX（GET：便于浏览器/axios blob 直接触发下载）。

    按大纲层级顺序组装全部 status=generated 且正文非空的章节；空章节跳过并在文末附清单。
    ``illustrations=true&stored_illustrations_only=true`` only reads persisted
    chapter images and embeds them. The legacy path without that flag is kept for
    older API clients, but the desktop export action never uses it.
    Worker J（净化层）：落库原文一律不动；导出文本统一过
    core.agent_engine.export_sanitizer.sanitize_export_text——剥离【n】引用锚点/
    硬门拒收原因（保留提取值）/插图建议变体/裸 JSON 泄漏，并按解读结果回填
    [采购人名称] 类占位符。配图装配在净化后的文本上进行（图片引用行不被二次净化）。
    """
    import os
    import tempfile
    from urllib.parse import quote

    from core.agent_engine.export_sanitizer import sanitize_export_text

    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    if force_illustrations:
        raise HTTPException(status_code=400, detail="导出不会现场生成配图，请先按章节生成并保存配图")

    chapter_result = await db.execute(select(Chapter).where(Chapter.project_id == project_id))
    chapters = chapter_result.scalars().all()
    chapters.sort(key=lambda c: (c.sort_order or 0, _chapter_natural_key(c.id)))

    # 导出净化（DB 原文不动，净化结果仅进内存 contents 覆盖表）
    fill_map = await _export_fill_map(db, project_id)
    contents: dict[str, str] = {}
    for c in chapters:
        raw = c.content or ""
        if not raw.strip():
            continue  # 空章保持"未生成清单"跳过语义（contents 中不存在 → 走原始判定）
        sanitized, _report = sanitize_export_text(raw, fill_map)
        contents[str(c.id)] = sanitized

    illustration_count = 0
    embedded_illustration_count = 0
    # Backward-compatible API path for older integrations that explicitly omit
    # stored_illustrations_only. The desktop workflow always sets it to true,
    # so its export action never invokes an image provider.
    if illustrations and not stored_illustrations_only:
        from types import SimpleNamespace

        view = [
            SimpleNamespace(id=c.id, content=contents.get(str(c.id), c.content or ""),
                            citation_ledger=c.citation_ledger)
            for c in chapters
            if (c.content or "").strip()
        ]
        assembled = await _assemble_export_illustrations(
            project_id,
            view,
            force_enabled=False,
            provider_override=illustration_provider,
            image_size=illustration_size,
            chapter_id=illustration_chapter_id,
        )
        contents.update(assembled)
        illustration_count = sum(text.count("![") for text in assembled.values())

    if illustrations:
        # Export is intentionally read-only for illustrations. Chapter-level
        # generation persists metadata first; this step only assembles those
        # saved references and never calls an image provider.
        from core.agent_engine.illustration_assembly import attach_stored_illustrations

        for c in chapters:
            ledger = c.citation_ledger if isinstance(c.citation_ledger, dict) else {}
            stored = ledger.get("_illustrations_generated") or []
            if not stored or not (c.content or "").strip():
                continue
            assembled = attach_stored_illustrations(contents.get(str(c.id), c.content or ""), stored)
            contents[str(c.id)] = assembled
            illustration_count += sum(1 for image in stored if isinstance(image, dict) and image.get("src"))

    docx_bytes, _skipped = _build_docx_from_chapters(project.name, chapters, contents or None)

    if contents:
        # Markdown 图片引用→图片对象嵌入（与排版管线 docx_format_skill 同一 helper）
        from core.agent_engine.illustration_assembly import embed_images_into_docx

        tmp_path = ""
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(docx_bytes)
                tmp_path = tmp.name
            embedded_illustration_count = embed_images_into_docx(tmp_path)
            with open(tmp_path, "rb") as f:
                docx_bytes = f.read()
        except Exception as e:  # noqa: BLE001
            logger.warning("直连导出 DOCX 配图嵌入失败（降级为引用文本）: %s", e)
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    filename = quote(f"{project.name}-正文直连导出.docx")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{filename}",
            "X-TenderPilot-Illustration-Count": str(embedded_illustration_count),
            "X-TenderPilot-Illustration-Requested": str(illustration_count),
            # Legacy aliases keep older desktop builds interoperable during rollout.
            "X-BidMaster-Illustration-Count": str(embedded_illustration_count),
            "X-BidMaster-Illustration-Requested": str(illustration_count),
        },
    )
